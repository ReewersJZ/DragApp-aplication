from tkinter import IntVar, PhotoImage, Checkbutton
from tkinter import Canvas
from docx.shared import Inches
import docx
import tkinter.messagebox
import webbrowser as wb
from PIL import Image, ImageDraw, ImageFont
import win32com.client as client
import datetime
import pymysql as mysql
import mysql.connector


# STYLE
wiel_czcionki = 12
kol_czcionki = 'white'
kol_tla = 'red'
wyglad = 'raised'
polozenie = 'e'
czcionka = 'Gill sans MT'
efekt_czcionki = 'bold'
kol_tla2 = 'grey'
wiel_czcionki2 = 9


# KLASY DANYCH:
class DaneFormularza:
    def __init__(self, data='brak', miasto='brak', ulica='brak', numer='brak', start_godzina='brak',
                 start_minuta='brak', koniec_godzina='brak', koniec_minuta='brak', stawka='brak', p1='brak', p2='brak',
                 p3='brak', p4='brak', p5='brak', p6='brak', p7='brak', p8='brak', p9='brak', p10='brak', p11='brak',
                 p12='brak', p13='brak', p14='brak', p15='brak', p16='brak', p17='brak', p18='brak', p19='brak',
                 p20='brak', p21='brak', p22='brak', p23='brak', p24='brak', p25='brak', p26='brak', p27='brak',
                 p28='brak', z1='brak', z2='brak', z3='brak', z4='brak', z5='brak', z6='brak', z7='brak', z8='brak',
                 z9='brak', z10='brak', z11='brak', z12='brak', z13='brak', z14='brak', z15='brak', z16='brak',
                 z17='brak', z18='brak', z19='brak', z20='brak', z21='brak', z22='brak', z23='brak', z24='brak',
                 z25='brak', z26='brak', z27='brak', z28='brak', status='nierozliczone', stawka_vat='brsk'):
        self.data = data
        self.miasto = miasto
        self.ulica = ulica
        self.numer = numer
        self.start_godzina = start_godzina
        self.start_minuta = start_minuta
        self.koniec_godzina = koniec_godzina
        self.koniec_minuta = koniec_minuta
        self.stawka = stawka
        # USED MATERIALS:
        self.p1, self.p2, self.p3, self.p4, self.p5, self.p6, self.p7, self.p8 = p1, p2, p3, p4, p5, p6, p7, p8
        self.p9, self.p10, self.p11, self.p12, self.p13, self.p14, self.p15 = p9, p10, p11, p12, p13, p14, p15
        self.p16, self.p17, self.p18, self.p19, self.p20, self.p21, self.p22 = p16, p17, p18, p19, p20, p21, p22
        self.p23, self.p24, self.p25, self.p26, self.p27, self.p28 = p23, p24, p25, p26, p27, p28
        # MEASUREMENT:
        self.z1, self.z2, self.z3, self.z4, self.z5, self.z6, self.z7, self.z8 = z1, z2, z3, z4, z5, z6, z7, z8
        self.z9, self.z10, self.z11, self.z12, self.z13, self.z14, self.z15 = z9, z10, z11, z12, z13, z14, z15
        self.z16, self.z17, self.z18, self.z19, self.z20, self.z21, self.z22 = z16, z17, z18, z19, z20, z21, z22
        self.z23, self.z24, self.z25, self.z26, self.z27, self.z28 = z23, z24, z25, z26, z27, z28
        self.status = status
        self.stawka_vat = stawka_vat

    def __str__(self):
        return f'{self.data}, {self.miasto}, {self.ulica}, {self.numer}'


class DaneFormularzaRozliczone:
    def __init__(self, data='brak', miasto='brak', ulica='brak', numer='brak', il_godzin='brak', stawka='brak',
                 status='nierozliczone', kwota_netto='brak', stawka_vat='brak', kwota='brak'):
        self.data = data
        self.miasto = miasto
        self.ulica = ulica
        self.numer = numer
        self.il_godzin = il_godzin
        self.stawka = stawka
        self.status = status
        self.kwota_netto = kwota_netto
        self.stawka_vat = stawka_vat
        self.kwota = kwota


class DaneFormularzaKalendarz:
    def __init__(self, id, data, miasto, ulica, numer, telefon, status, uwagi):
        self.id = id
        self.data = data
        self.miasto = miasto
        self.ulica = ulica
        self.numer = numer
        self.telefon = telefon
        self.status = status
        self.uwagi = uwagi


def main():
    global poleTekstowe, poleTekstowe2, poleTekstowe3, poleTekstowe6, koncowy_wynik, root
    root = tkinter.Tk()
    root.title('DragApp')
    root.geometry('1200x900')
    frame = tkinter.Frame(root, bg='white')
    frame.place(relx=0.025, rely=0.025, relwidth=0.95, relheight=0.95)

    zdjecie = tkinter.PhotoImage(file='grafika\\dach.png\\')
    zdjecie1 = tkinter.Label(master=frame, image=zdjecie, bg='white')
    zdjecie1.grid()

    button1 = tkinter.Button(master=frame, text='NOWY\n FORMULARZ', bg='grey', fg='white', command=menu_pierwszy,
                             relief='raised', anchor='center', font=('Gill sans MT', 12, 'bold'))
    button1.place(relx=0.6455, rely=0.1, relwidth=0.15, relheight=0.15)
    button2 = tkinter.Button(master=frame, text='POBIERZ\n Z KALENDARZA', bg='red', fg='white', command=menu_trzeci,
                             relief='raised', anchor='center', font=('Gill sans MT', 12, 'bold'))
    button2.place(relx=0.8, rely=0.1, relwidth=0.15, relheight=0.15)
    button3 = tkinter.Button(master=frame, text='RAPORTY', bg='#37474F', fg='white', command=menu_drugi,
                             relief='raised', anchor='center', font=('Gill sans MT', 12, 'bold'))
    button3.place(relx=0.6455, rely=0.2555, relwidth=0.15, relheight=0.15)
    button4 = tkinter.Button(master=frame, text='KALENDARZ', bg='#607D8B', fg='white', command=menu_czwarty,
                             relief='raised', anchor='center', font=('Gill sans MT', 12, 'bold'))
    button4.place(relx=0.8, rely=0.2555, relwidth=0.15, relheight=0.15)
    opis_programu = tkinter.Label(master=frame, text='DragApp v1.1. 2020 Wszelkie prawa zastrzeżone. '
                                                     'Projekt i realizacja REEWERS Justyna Zahraj', fg='#939393',
                                  bg='white', anchor='center', font=(czcionka, 10, efekt_czcionki))
    opis_programu.place(relx=0.05, rely=0.9, relwidth=0.9, relheight=0.05)
    logo = tkinter.PhotoImage(file='grafika\\logo1.png')
    foto = tkinter.Label(master=root, image=logo, bg='white')
    foto.place(relx=0.08, rely=0.85, relwidth=0.075, relheight=0.075)

    root.mainloop()


def menu_pierwszy():
    global app_a, poleTekstowe, poleTekstowe2, poleTekstowe3, poleTekstowe4, poleTekstowe5, poleTekstowe6, poleTekstowe7
    global poleTekstowe8, poleTekstowe9, poleTekstowe10, wybrana_stawka, tekst_produkt_1, tekst_produkt_2
    global tekst_produkt_3, tekst_produkt_4, tekst_produkt_5, tekst_produkt_6, tekst_produkt_7, tekst_produkt_8
    global tekst_produkt_9, tekst_produkt_10, tekst_produkt_11, tekst_produkt_12, tekst_produkt_13, tekst_produkt_14
    global tekst_produkt_15, tekst_produkt_16, tekst_produkt_17, tekst_produkt_18, tekst_produkt_19, tekst_produkt_20
    global tekst_produkt_21, tekst_produkt_22, tekst_produkt_23, tekst_produkt_24, tekst_produkt_25, tekst_produkt_26
    global tekst_produkt_27, tekst_produkt_28, root, wybrana_stawka, koncowy_wynik, pomiar_kolejny
    global tabela_pomocnicza_pomiary, pomoc_pomiary1, pomoc_pomiary2, pomoc_pomiary3, pomoc_pomiary4, pomoc_pomiary5
    global pomoc_pomiary6, pomoc_pomiary7, pomoc_pomiary8, pomoc_pomiary9, pomoc_pomiary10, pomoc_pomiary11
    global pomoc_pomiary12, pomoc_pomiary13, pomoc_pomiary14, pomoc_pomiary15, pomoc_pomiary16, pomoc_pomiary17
    global pomoc_pomiary18, pomoc_pomiary19, pomoc_pomiary20, pomoc_pomiary21, pomoc_pomiary22, pomoc_pomiary23
    global pomoc_pomiary24, pomoc_pomiary25, pomoc_pomiary26, pomoc_pomiary27, pomoc_pomiary28
    global stawka_vat_value, stawka_value

    root_a = tkinter.Toplevel()
    root_a.title('DragApp')
    root_a.geometry('1200x900')
    app_a = tkinter.Frame(root_a, bg='white')
    app_a.place(relx=0.025, rely=0.025, relwidth=0.95, relheight=0.95)

    logo = tkinter.PhotoImage(file='grafika\\logo.png')
    foto = tkinter.Label(master=app_a, image=logo, bg='white')
    foto.place(relx=0.7, rely=0.1, relwidth=0.3, relheight=0.2)

    # TABELA NAGŁÓWKOWA:
    label = tkinter.Label(master=app_a, text='DATA  ', fg=kol_czcionki, bg=kol_tla, relief=wyglad, anchor=polozenie,
                          font=(czcionka, wiel_czcionki, efekt_czcionki))
    label.place(relx=0.1, rely=0.1, relwidth=0.1, relheight=0.05)
    poleTekstowe = tkinter.Entry(app_a, bg='white', font=('Gill sans MT', 14), justify='center')
    poleTekstowe.place(relx=0.2, rely=0.1, relwidth=0.2, relheight=0.05)

    label2 = tkinter.Label(master=app_a, text='MIASTO  ', fg=kol_czcionki, bg=kol_tla, relief=wyglad, anchor=polozenie,
                           font=(czcionka, wiel_czcionki, efekt_czcionki))
    label2.place(relx=0.4, rely=0.1, relwidth=0.1, relheight=0.05)
    poleTekstowe2 = tkinter.Entry(app_a, bg='white', font=('Gill sans MT', 14), justify='center')
    poleTekstowe2.place(relx=0.5, rely=0.1, relwidth=0.2, relheight=0.05)

    label3 = tkinter.Label(master=app_a, text='ULICA  ', fg=kol_czcionki, bg=kol_tla, relief=wyglad, anchor=polozenie,
                           font=(czcionka, wiel_czcionki, efekt_czcionki))
    label3.place(relx=0.1, rely=0.175, relwidth=0.1, relheight=0.05)
    poleTekstowe3 = tkinter.Entry(app_a, bg='white', font=('Gill sans MT', 14), justify='center')
    poleTekstowe3.place(relx=0.2, rely=0.175, relwidth=0.2, relheight=0.05)

    label4 = tkinter.Label(master=app_a, text='NUMER  ', fg=kol_czcionki, bg=kol_tla, relief=wyglad, anchor=polozenie,
                           font=(czcionka, wiel_czcionki, efekt_czcionki))
    label4.place(relx=0.4, rely=0.175, relwidth=0.1, relheight=0.05)
    poleTekstowe4 = tkinter.Entry(app_a, bg='white', font=('Gill sans MT', 14), justify='center')
    poleTekstowe4.place(relx=0.5, rely=0.175, relwidth=0.2, relheight=0.05)

    label5 = tkinter.Label(master=app_a, text='START', fg=kol_czcionki, bg='#37474F', relief=wyglad, anchor='center',
                           font=(czcionka, wiel_czcionki, efekt_czcionki))
    label5.place(relx=0.1, rely=0.25, relwidth=0.3, relheight=0.05)
    label6 = tkinter.Label(master=app_a, text='godz.  ', fg=kol_czcionki, bg=kol_tla, relief=wyglad, anchor=polozenie,
                           font=(czcionka, wiel_czcionki, efekt_czcionki))
    label6.place(relx=0.1, rely=0.3, relwidth=0.1, relheight=0.05)
    poleTekstowe6 = tkinter.Entry(app_a, bg='white', font=('Gill sans MT', 14), justify='center')
    poleTekstowe6.place(relx=0.2, rely=0.3, relwidth=0.05, relheight=0.05)
    label7 = tkinter.Label(master=app_a, text='min.  ', fg=kol_czcionki, bg=kol_tla, relief=wyglad, anchor=polozenie,
                           font=(czcionka, wiel_czcionki, efekt_czcionki))
    label7.place(relx=0.25, rely=0.3, relwidth=0.1, relheight=0.05)
    poleTekstowe7 = tkinter.Entry(app_a, bg='white', font=('Gill sans MT', 14), justify='center')
    poleTekstowe7.place(relx=0.35, rely=0.3, relwidth=0.05, relheight=0.05)

    label8 = tkinter.Label(master=app_a, text='KONIEC', fg=kol_czcionki, bg='#37474F', relief=wyglad, anchor='center',
                           font=(czcionka, wiel_czcionki, efekt_czcionki))
    label8.place(relx=0.1, rely=0.375, relwidth=0.3, relheight=0.05)
    label9 = tkinter.Label(master=app_a, text='godz.  ', fg=kol_czcionki, bg=kol_tla, relief=wyglad, anchor=polozenie,
                           font=(czcionka, wiel_czcionki, efekt_czcionki))
    label9.place(relx=0.1, rely=0.425, relwidth=0.1, relheight=0.05)
    poleTekstowe9 = tkinter.Entry(app_a, bg='white', width=30, font=('Gill sans MT', 14), justify='center')
    poleTekstowe9.place(relx=0.2, rely=0.425, relwidth=0.05, relheight=0.05)
    label10 = tkinter.Label(master=app_a, text='min.  ', fg=kol_czcionki, bg=kol_tla, relief=wyglad, anchor=polozenie,
                            font=(czcionka, wiel_czcionki, efekt_czcionki))
    label10.place(relx=0.25, rely=0.425, relwidth=0.1, relheight=0.05)
    poleTekstowe10 = tkinter.Entry(app_a, bg='white', font=('Gill sans MT', 14), justify='center')
    poleTekstowe10.place(relx=0.35, rely=0.425, relwidth=0.05, relheight=0.05)

    label11 = tkinter.Label(master=app_a, text='STAWKA / VAT', fg=kol_czcionki, bg='#37474F', relief=wyglad,
                            anchor='center',
                            font=(czcionka, wiel_czcionki, efekt_czcionki))
    label11.place(relx=0.5, rely=0.3, relwidth=0.2, relheight=0.05)
    stawka_value = IntVar()
    stawka1 = tkinter.Radiobutton(master=app_a, text="2 zł", value=2, variable=stawka_value, tristatevalue=1,
                                  activebackground='red', anchor='w', font=(czcionka, wiel_czcionki, efekt_czcionki))
    stawka1.place(relx=0.5, rely=0.375, relwidth=0.095, relheight=0.05)
    stawka2 = tkinter.Radiobutton(master=app_a, text="4 zł", value=4, variable=stawka_value, tristatevalue=1,
                                  activebackground='red', anchor='w', font=(czcionka, wiel_czcionki, efekt_czcionki))
    stawka2.place(relx=0.5, rely=0.425, relwidth=0.095, relheight=0.05)

    stawka_vat_value = IntVar()
    stawka_vat1 = tkinter.Radiobutton(master=app_a, text="8 %", value=8, variable=stawka_vat_value, tristatevalue=1,
                                      activebackground='red', anchor='w',
                                      font=(czcionka, wiel_czcionki, efekt_czcionki))
    stawka_vat1.place(relx=0.6, rely=0.375, relwidth=0.1, relheight=0.05)
    stawka_vat2 = tkinter.Radiobutton(master=app_a, text="23 %", value=23, variable=stawka_vat_value, tristatevalue=1,
                                      activebackground='red', anchor='w',
                                      font=(czcionka, wiel_czcionki, efekt_czcionki))
    stawka_vat2.place(relx=0.6, rely=0.425, relwidth=0.1, relheight=0.05)

    # TABELA Z MATERIAŁAMI:
    label12 = tkinter.Label(master=app_a, text='UŻYTE MATERIAŁY', fg=kol_czcionki, bg='#37474F', relief=wyglad,
                            anchor='center', font=(czcionka, wiel_czcionki2, efekt_czcionki))
    label12.place(relx=0.1, rely=0.525, relwidth=0.6, relheight=0.05)
    list_of_materials = ['UZIOMY ', 'ZŁĄCZKI 2 ŚRUB. ', 'ZŁĄCZKI 4 ŚRUB. ', 'ZŁĄCZKI 4 ŚRUB. \nBEDNARKA ']
    create_labels(list_of_materials, 0.1, 0.575)
    list_of_materials2 = ['ZŁĄCZKI 45° \n2 ŚRUB. ', 'ZŁĄCZE KONTR. ', 'PUSZKI PODT. ', 'PUSZKI GRUNT. ']
    create_labels(list_of_materials2, 0.1, 0.625)
    list_of_materials3 = ['T-ki ', 'L-ki ', 'GĄSIOR METAL ', 'GĄSIOR \nMALOWANY ']
    create_labels(list_of_materials3, 0.1, 0.675)
    list_of_materials4 = ['NACIĄG 20 CM ', 'KOTWY 18 ', 'KOTWY 20 ', 'USZCZELNIACZ ']
    create_labels(list_of_materials4, 0.1, 0.725)
    list_of_materials5 = ['BETONIKI ', 'KLEJ ', 'DRUT ALU. ', 'DRUT STAL ']
    create_labels(list_of_materials5, 0.1, 0.775)
    list_of_materials6 = ['POLWINIT BIAŁY ', 'POLWINIT \nCZARNY ', 'BEDNARKA 30x4 ', 'BEDNARKA 25x4 ']
    create_labels(list_of_materials6, 0.1, 0.825)
    list_of_materials7 = ['SZTYCA 4m ', 'INNE']
    create_labels(list_of_materials7, 0.1, 0.875)

    tekst_produkt_1 = tkinter.Entry(app_a, bg='white', width=30, font=('Gill sans MT', 14), justify='center')
    tekst_produkt_1.place(relx=0.2, rely=0.575, relwidth=0.05, relheight=0.05)
    tekst_produkt_2 = tkinter.Entry(app_a, bg='white', font=('Gill sans MT', 14), justify='center')
    tekst_produkt_2.place(relx=0.35, rely=0.575, relwidth=0.05, relheight=0.05)
    tekst_produkt_3 = tkinter.Entry(app_a, bg='white', width=30, font=('Gill sans MT', 14), justify='center')
    tekst_produkt_3.place(relx=0.5, rely=0.575, relwidth=0.05, relheight=0.05)
    tekst_produkt_4 = tkinter.Entry(app_a, bg='white', font=('Gill sans MT', 14), justify='center')
    tekst_produkt_4.place(relx=0.65, rely=0.575, relwidth=0.05, relheight=0.05)

    tekst_produkt_5 = tkinter.Entry(app_a, bg='white', width=30, font=('Gill sans MT', 14), justify='center')
    tekst_produkt_5.place(relx=0.2, rely=0.625, relwidth=0.05, relheight=0.05)
    tekst_produkt_6 = tkinter.Entry(app_a, bg='white', font=('Gill sans MT', 14), justify='center')
    tekst_produkt_6.place(relx=0.35, rely=0.625, relwidth=0.05, relheight=0.05)
    tekst_produkt_7 = tkinter.Entry(app_a, bg='white', width=30, font=('Gill sans MT', 14), justify='center')
    tekst_produkt_7.place(relx=0.5, rely=0.625, relwidth=0.05, relheight=0.05)
    tekst_produkt_8 = tkinter.Entry(app_a, bg='white', font=('Gill sans MT', 14), justify='center')
    tekst_produkt_8.place(relx=0.65, rely=0.625, relwidth=0.05, relheight=0.05)

    tekst_produkt_9 = tkinter.Entry(app_a, bg='white', width=30, font=('Gill sans MT', 14), justify='center')
    tekst_produkt_9.place(relx=0.2, rely=0.675, relwidth=0.05, relheight=0.05)
    tekst_produkt_10 = tkinter.Entry(app_a, bg='white', font=('Gill sans MT', 14), justify='center')
    tekst_produkt_10.place(relx=0.35, rely=0.675, relwidth=0.05, relheight=0.05)
    tekst_produkt_11 = tkinter.Entry(app_a, bg='white', width=30, font=('Gill sans MT', 14), justify='center')
    tekst_produkt_11.place(relx=0.5, rely=0.675, relwidth=0.05, relheight=0.05)
    tekst_produkt_12 = tkinter.Entry(app_a, bg='white', font=('Gill sans MT', 14), justify='center')
    tekst_produkt_12.place(relx=0.65, rely=0.675, relwidth=0.05, relheight=0.05)

    tekst_produkt_13 = tkinter.Entry(app_a, bg='white', width=30, font=('Gill sans MT', 14), justify='center')
    tekst_produkt_13.place(relx=0.2, rely=0.725, relwidth=0.05, relheight=0.05)
    tekst_produkt_14 = tkinter.Entry(app_a, bg='white', font=('Gill sans MT', 14), justify='center')
    tekst_produkt_14.place(relx=0.35, rely=0.725, relwidth=0.05, relheight=0.05)
    tekst_produkt_15 = tkinter.Entry(app_a, bg='white', width=30, font=('Gill sans MT', 14), justify='center')
    tekst_produkt_15.place(relx=0.5, rely=0.725, relwidth=0.05, relheight=0.05)
    tekst_produkt_16 = tkinter.Entry(app_a, bg='white', font=('Gill sans MT', 14), justify='center')
    tekst_produkt_16.place(relx=0.65, rely=0.725, relwidth=0.05, relheight=0.05)

    tekst_produkt_17 = tkinter.Entry(app_a, bg='white', font=('Gill sans MT', 14), justify='center')
    tekst_produkt_17.place(relx=0.2, rely=0.775, relwidth=0.05, relheight=0.05)
    tekst_produkt_18 = tkinter.Entry(app_a, bg='white', font=('Gill sans MT', 14), justify='center')
    tekst_produkt_18.place(relx=0.35, rely=0.775, relwidth=0.05, relheight=0.05)
    tekst_produkt_19 = tkinter.Entry(app_a, bg='white', font=('Gill sans MT', 14), justify='center')
    tekst_produkt_19.place(relx=0.5, rely=0.775, relwidth=0.05, relheight=0.05)
    tekst_produkt_20 = tkinter.Entry(app_a, bg='white', font=('Gill sans MT', 14), justify='center')
    tekst_produkt_20.place(relx=0.65, rely=0.775, relwidth=0.05, relheight=0.05)

    tekst_produkt_21 = tkinter.Entry(app_a, bg='white', font=('Gill sans MT', 14), justify='center')
    tekst_produkt_21.place(relx=0.2, rely=0.825, relwidth=0.05, relheight=0.05)
    tekst_produkt_22 = tkinter.Entry(app_a, bg='white', font=('Gill sans MT', 14), justify='center')
    tekst_produkt_22.place(relx=0.35, rely=0.825, relwidth=0.05, relheight=0.05)
    tekst_produkt_23 = tkinter.Entry(app_a, bg='white', font=('Gill sans MT', 14), justify='center')
    tekst_produkt_23.place(relx=0.5, rely=0.825, relwidth=0.05, relheight=0.05)
    tekst_produkt_24 = tkinter.Entry(app_a, bg='white', font=('Gill sans MT', 14), justify='center')
    tekst_produkt_24.place(relx=0.65, rely=0.825, relwidth=0.05, relheight=0.05)

    tekst_produkt_25 = tkinter.Entry(app_a, bg='white', font=('Gill sans MT', 14), justify='center')
    tekst_produkt_25.place(relx=0.2, rely=0.875, relwidth=0.05, relheight=0.05)
    tekst_produkt_26 = tkinter.Entry(app_a, bg='white', font=('Gill sans MT', 10), justify='center')
    tekst_produkt_26.place(relx=0.35, rely=0.875, relwidth=0.35, relheight=0.05)
    tekst_produkt_27 = tkinter.Entry(app_a, bg='white', font=('Gill sans MT', 14), justify='center')
    tekst_produkt_27.forget()
    # (relx=0.5, rely=0.875, relwidth=0.05, relheight=0.05) - dodatkowa pozycja w materialach
    tekst_produkt_28 = tkinter.Entry(app_a, bg='white', font=('Gill sans MT', 14), justify='center')
    tekst_produkt_28.forget()
    # (relx=0.65, rely=0.875, relwidth=0.05, relheight=0.05) - dodatkowa pozycja w materialach (gdyby
    # trzeba bylo przywrócić to należy dodać opis pozycji w def create_labels

    button2 = tkinter.Button(master=app_a, text="DODAJ POMIARY", bg=kol_tla2, fg='white',
                             font=(czcionka, wiel_czcionki, efekt_czcionki), command=pobierz_pomiary)
    button2.place(relx=0.75, rely=0.615, relwidth=0.2, relheight=0.05)
    button3 = tkinter.Button(master=app_a, text="DODAJ RYSUNEK", bg='#37474F', fg='white',
                             font=(czcionka, wiel_czcionki, efekt_czcionki), command=pobierz_rysunek)
    button3.place(relx=0.75, rely=0.67, relwidth=0.2, relheight=0.05)
    button1 = tkinter.Button(master=app_a, text="ZAPISZ", bg='red', fg='white',
                             font=(czcionka, wiel_czcionki, efekt_czcionki), command=akcja)
    button1.place(relx=0.75, rely=0.725, relwidth=0.2, relheight=0.05)
    koncowy_wynik = tkinter.Label(master=app_a, text='', fg='black', bg='white',
                                  font=(czcionka, wiel_czcionki, efekt_czcionki), anchor='center')
    koncowy_wynik.place(relx=0.75, rely=0.8, relwidth=0.2, relheight=0.05)

    # TABELA POMOCNICZA
    pomoc_pomiary1 = tkinter.Label(master=app_a, text='')
    pomoc_pomiary1.place_forget()
    pomoc_pomiary2 = tkinter.Label(master=app_a, text='')
    pomoc_pomiary2.place_forget()
    pomoc_pomiary3 = tkinter.Label(master=app_a, text='')
    pomoc_pomiary3.place_forget()
    pomoc_pomiary4 = tkinter.Label(master=app_a, text='')
    pomoc_pomiary4.place_forget()
    pomoc_pomiary5 = tkinter.Label(master=app_a, text='')
    pomoc_pomiary5.place_forget()
    pomoc_pomiary6 = tkinter.Label(master=app_a, text='')
    pomoc_pomiary6.place_forget()
    pomoc_pomiary7 = tkinter.Label(master=app_a, text='')
    pomoc_pomiary7.place_forget()
    pomoc_pomiary8 = tkinter.Label(master=app_a, text='')
    pomoc_pomiary8.place_forget()
    pomoc_pomiary9 = tkinter.Label(master=app_a, text='')
    pomoc_pomiary9.place_forget()
    pomoc_pomiary10 = tkinter.Label(master=app_a, text='')
    pomoc_pomiary10.place_forget()
    pomoc_pomiary11 = tkinter.Label(master=app_a, text='')
    pomoc_pomiary11.place_forget()
    pomoc_pomiary12 = tkinter.Label(master=app_a, text='')
    pomoc_pomiary12.place_forget()
    pomoc_pomiary13 = tkinter.Label(master=app_a, text='')
    pomoc_pomiary13.place_forget()
    pomoc_pomiary14 = tkinter.Label(master=app_a, text='')
    pomoc_pomiary14.place_forget()
    pomoc_pomiary15 = tkinter.Label(master=app_a, text='')
    pomoc_pomiary15.place_forget()
    pomoc_pomiary16 = tkinter.Label(master=app_a, text='')
    pomoc_pomiary16.place_forget()
    pomoc_pomiary17 = tkinter.Label(master=app_a, text='')
    pomoc_pomiary17.place_forget()
    pomoc_pomiary18 = tkinter.Label(master=app_a, text='')
    pomoc_pomiary18.place_forget()
    pomoc_pomiary19 = tkinter.Label(master=app_a, text='')
    pomoc_pomiary19.place_forget()
    pomoc_pomiary20 = tkinter.Label(master=app_a, text='')
    pomoc_pomiary20.place_forget()
    pomoc_pomiary21 = tkinter.Label(master=app_a, text='')
    pomoc_pomiary21.place_forget()
    pomoc_pomiary22 = tkinter.Label(master=app_a, text='')
    pomoc_pomiary22.place_forget()
    pomoc_pomiary23 = tkinter.Label(master=app_a, text='')
    pomoc_pomiary23.place_forget()
    pomoc_pomiary24 = tkinter.Label(master=app_a, text='')
    pomoc_pomiary24.place_forget()
    pomoc_pomiary25 = tkinter.Label(master=app_a, text='')
    pomoc_pomiary25.place_forget()
    pomoc_pomiary26 = tkinter.Label(master=app_a, text='')
    pomoc_pomiary26.place_forget()
    pomoc_pomiary27 = tkinter.Label(master=app_a, text='')
    pomoc_pomiary27.place_forget()
    pomoc_pomiary28 = tkinter.Label(master=app_a, text='')
    pomoc_pomiary28.place_forget()
    tabela_pomocnicza_pomiary = tkinter.Label(master=app_a, text='')
    tabela_pomocnicza_pomiary.place_forget()

    root.mainloop()


def menu_drugi():
    global pole_data_od, pole_data_do, root_drugi
    global do_rozliczenia1, do_rozliczenia2, do_rozliczenia3, do_rozliczenia4, do_rozliczenia5, do_rozliczenia6
    global do_rozliczenia7, do_rozliczenia8, do_rozliczenia9, do_rozliczenia10, do_rozliczenia11, do_rozliczenia12
    global do_rozliczenia13, do_rozliczenia14
    global lista_var

    root_drugi = tkinter.Toplevel()
    root_drugi.title('DragApp')
    root_drugi.geometry('1200x900')
    app_drugi = tkinter.Frame(root_drugi, bg='white')
    app_drugi.place(relx=0.025, rely=0.025, relwidth=0.95, relheight=0.95)

    logo = tkinter.PhotoImage(file='grafika\\logo.png')
    foto = tkinter.Label(master=app_drugi, image=logo, bg='white')
    foto.place(relx=0.7, rely=0.1, relwidth=0.3, relheight=0.2)

    label_data_od = tkinter.Label(master=app_drugi, text='DATA OD:  ', fg=kol_czcionki, bg=kol_tla, relief=wyglad,
                                  anchor=polozenie, font=(czcionka, wiel_czcionki, efekt_czcionki))
    label_data_od.place(relx=0.1, rely=0.1, relwidth=0.1, relheight=0.05)
    pole_data_od = tkinter.Entry(app_drugi, bg='white', font=('Gill sans MT', 14), justify='center')
    pole_data_od.place(relx=0.2, rely=0.1, relwidth=0.2, relheight=0.05)

    label_data_do = tkinter.Label(master=app_drugi, text='DATA DO:  ', fg=kol_czcionki, bg=kol_tla, relief=wyglad,
                                  anchor=polozenie, font=(czcionka, wiel_czcionki, efekt_czcionki))
    label_data_do.place(relx=0.4, rely=0.1, relwidth=0.1, relheight=0.05)
    pole_data_do = tkinter.Entry(app_drugi, bg='white', font=('Gill sans MT', 14), justify='center')
    pole_data_do.place(relx=0.5, rely=0.1, relwidth=0.2, relheight=0.05)

    przycisk_zatwierdz = tkinter.Button(app_drugi, text='ZATWIERDŹ', bg='#37474F', fg='white',
                                        font=(czcionka, wiel_czcionki, efekt_czcionki), command=zatwierdz_przedzial_dat)
    przycisk_zatwierdz.place(relx=0.5, rely=0.175, relwidth=0.2, relheight=0.05)

    przycisk_zatwierdz_do_rozliczenia = tkinter.Button(app_drugi, text='GENERUJ RAPORT', bg='#37474F', fg='white',
                                                       font=(czcionka, wiel_czcionki, efekt_czcionki),
                                                       command=var_states)
    przycisk_zatwierdz_do_rozliczenia.place(relx=0.75, rely=0.67, relwidth=0.2, relheight=0.05)

    var1 = IntVar()
    do_rozliczenia1 = Checkbutton(app_drugi, text='', variable=var1)
    do_rozliczenia1.forget()
    var2 = IntVar()
    do_rozliczenia2 = Checkbutton(app_drugi, text='', variable=var2)
    do_rozliczenia2.forget()
    var3 = IntVar()
    do_rozliczenia3 = Checkbutton(app_drugi, text='', variable=var3)
    do_rozliczenia3.forget()
    var4 = IntVar()
    do_rozliczenia4 = Checkbutton(app_drugi, text='', variable=var4)
    do_rozliczenia4.forget()
    var5 = IntVar()
    do_rozliczenia5 = Checkbutton(app_drugi, text='', variable=var5)
    do_rozliczenia5.forget()
    var6 = IntVar()
    do_rozliczenia6 = Checkbutton(app_drugi, text='', variable=var6)
    do_rozliczenia6.forget()
    var7 = IntVar()
    do_rozliczenia7 = Checkbutton(app_drugi, text='', variable=var7)
    do_rozliczenia7.forget()
    var8 = IntVar()
    do_rozliczenia8 = Checkbutton(app_drugi, text='', variable=var8)
    do_rozliczenia8.forget()
    var9 = IntVar()
    do_rozliczenia9 = Checkbutton(app_drugi, text='', variable=var9)
    do_rozliczenia9.forget()
    var10 = IntVar()
    do_rozliczenia10 = Checkbutton(app_drugi, text='', variable=var10)
    do_rozliczenia10.forget()
    var11 = IntVar()
    do_rozliczenia11 = Checkbutton(app_drugi, text='', variable=var11)
    do_rozliczenia11.forget()
    var12 = IntVar()
    do_rozliczenia12 = Checkbutton(app_drugi, text='', variable=var12)
    do_rozliczenia12.forget()
    var13 = IntVar()
    do_rozliczenia13 = Checkbutton(app_drugi, text='', variable=var13)
    do_rozliczenia13.forget()
    var14 = IntVar()
    do_rozliczenia14 = Checkbutton(app_drugi, text='', variable=var14)
    do_rozliczenia14.forget()

    lista_var = [var1, var2, var3, var4, var5, var6, var7, var8, var9, var10, var11, var12, var13, var14]

    root_drugi.mainloop()


def menu_trzeci(): # W FAZIE PROJEKTOWANIA
    global root, lista_checkboxow_kalendarz, lista_do_kalendarza, lista_var_kalendarz, e
    pass
    """a = 0
    zbior_rozliczonych_do_worda = []

    for element in lista_checkboxow_kalendarz:
        pobrane = lista_var_kalendarz[a].get()
        if pobrane == 1:

            print(pobrane)
        a += 1

    e = tkinter.Entry(root, width=10)
    e.pack()

    b1 = tkinter.Button(root, text="animal", command=lambda: set_text(data, miasto, ulica, numer))
    b1.pack()

    b2 = tkinter.Button(root, text="plant", command=lambda: set_text("plant"))
    b2.pack()

    def set_text(text):
        global e
        e.delete(0, 'end')
        e.insert(0, text)
        return"""


def menu_czwarty():
    global lista_checkboxow_kalendarz, lista_var_kalendarz, lista_do_kalendarza

    a = mysql.connector.connect(host='127.0.0.1', database='dragapp', user='root', password='')

    lista_do_kalendarza = []
    lista_do_wyswietlenia = []

    kursor = a.cursor()
    pytanie = 'SELECT * FROM kalendarz WHERE status="nowy"'
    kursor.execute(pytanie)
    for element in kursor:
        id = element[0]
        data = element[1]
        data = str(data.day) + '-' + str(data.month) + '-' + str(data.year)
        miasto = element[2]
        ulica = element[3]
        numer = element[4]
        telefon = element[5]
        status = element[6]
        uwagi = element[7]
        nowy = DaneFormularzaKalendarz(id, data, miasto, ulica, numer, telefon, status, uwagi)
        lista_do_kalendarza.append(nowy)

        lista_do_wyswietlenia.append(f'{nowy.data}   {nowy.miasto}, ul. {nowy.ulica} {nowy.numer}        '
                                     f'/ tel. {nowy.telefon} /')

    kalendarz_var1 = IntVar()
    kalendarz_poz1 = Checkbutton(root, text='', variable=kalendarz_var1)
    kalendarz_poz1.forget()
    kalendarz_var2 = IntVar()
    kalendarz_poz2 = Checkbutton(root, text='', variable=kalendarz_var2)
    kalendarz_poz2.forget()
    kalendarz_var3 = IntVar()
    kalendarz_poz3 = Checkbutton(root, text='', variable=kalendarz_var3)
    kalendarz_poz3.forget()
    kalendarz_var4 = IntVar()
    kalendarz_poz4 = Checkbutton(root, text='', variable=kalendarz_var4)
    kalendarz_poz4.forget()
    kalendarz_var5 = IntVar()
    kalendarz_poz5 = Checkbutton(root, text='', variable=kalendarz_var5)
    kalendarz_poz5.forget()
    kalendarz_var6 = IntVar()
    kalendarz_poz6 = Checkbutton(root, text='', variable=kalendarz_var6)
    kalendarz_poz6.forget()
    kalendarz_var7 = IntVar()
    kalendarz_poz7 = Checkbutton(root, text='', variable=kalendarz_var7)
    kalendarz_poz7.forget()

    lista_var_kalendarz = [kalendarz_var1, kalendarz_var2, kalendarz_var3, kalendarz_var4, kalendarz_var5,
                           kalendarz_var6, kalendarz_var7]
    lista_checkboxow_kalendarz = [kalendarz_poz1, kalendarz_poz2, kalendarz_poz3, kalendarz_poz4, kalendarz_poz5,
                                  kalendarz_poz6, kalendarz_poz7]

    tworz_checkbutton_do_rozliczen(lista_do_wyswietlenia, lista_checkboxow_kalendarz, 0.25, 0.5, 'white', 'black')


def akcja():
    global poleTekstowe, poleTekstowe2, poleTekstowe3, poleTekstowe4, poleTekstowe5, poleTekstowe6, poleTekstowe7
    global poleTekstowe8, poleTekstowe9, poleTekstowe10, wybrana_stawka, tekst_produkt_1, tekst_produkt_2
    global tekst_produkt_3, tekst_produkt_4, tekst_produkt_5, tekst_produkt_6, tekst_produkt_7, tekst_produkt_8
    global tekst_produkt_9, tekst_produkt_10, tekst_produkt_11, tekst_produkt_12, tekst_produkt_13, tekst_produkt_14
    global tekst_produkt_15, tekst_produkt_16, tekst_produkt_17, tekst_produkt_18, tekst_produkt_19, tekst_produkt_20
    global tekst_produkt_21, tekst_produkt_22, tekst_produkt_23, tekst_produkt_24, tekst_produkt_25, tekst_produkt_26
    global tekst_produkt_27, tekst_produkt_28, root, wybrana_stawka, koncowy_wynik
    global pomiar_kolejny, tabela_pomiary, wszystkie_pomiary, lista_z_pomiarami
    global pierwszy_pomiar_wejcie, pierwszy_pomiar_wejcie1
    global kolejny_wejcie, kolejny_wejscie1, kolejny_wejscie2, kolejny_wejscie3
    global kolejny_wejscie4, kolejny_wejscie5, kolejny_wejscie6, kolejny_wejscie7, kolejny_wejscie8, kolejny_wejscie9
    global kolejny_wejscie10, kolejny_wejscie11, kolejny_wejscie12, kolejny_wejscie13, kolejny_wejscie14
    global kolejny_wejscie15, kolejny_wejscie16, kolejny_wejscie17, kolejny_wejscie18, kolejny_wejscie19
    global kolejny_wejscie20, kolejny_wejscie21, kolejny_wejscie22, kolejny_wejscie23, kolejny_wejscie24
    global kolejny_wejscie25, stawka_vat_value, stawka_value

    data = poleTekstowe.get()
    miasto = poleTekstowe2.get()
    ulica = poleTekstowe3.get()
    numer = poleTekstowe4.get()
    start_godzina = poleTekstowe6.get()
    start_minuta = poleTekstowe7.get()
    koniec_godzina = poleTekstowe9.get()
    koniec_minuta = poleTekstowe10.get()
    stawka = stawka_value.get()
    p1 = tekst_produkt_1.get()
    p2 = tekst_produkt_2.get()
    p3 = tekst_produkt_3.get()
    p4 = tekst_produkt_4.get()
    p5 = tekst_produkt_5.get()
    p6 = tekst_produkt_6.get()
    p7 = tekst_produkt_7.get()
    p8 = tekst_produkt_8.get()
    p9 = tekst_produkt_9.get()
    p10 = tekst_produkt_10.get()
    p11 = tekst_produkt_11.get()
    p12 = tekst_produkt_12.get()
    p13 = tekst_produkt_13.get()
    p14 = tekst_produkt_14.get()
    p15 = tekst_produkt_15.get()
    p16 = tekst_produkt_16.get()
    p17 = tekst_produkt_17.get()
    p18 = tekst_produkt_18.get()
    p19 = tekst_produkt_19.get()
    p20 = tekst_produkt_20.get()
    p21 = tekst_produkt_21.get()
    p22 = tekst_produkt_22.get()
    p23 = tekst_produkt_23.get()
    p24 = tekst_produkt_24.get()
    p25 = tekst_produkt_25.get()
    p26 = tekst_produkt_26.get()
    p27 = tekst_produkt_27.get()
    p28 = tekst_produkt_28.get()
    z1, z2, z3 = pomoc_pomiary1['text'], pomoc_pomiary2['text'], pomoc_pomiary3['text']
    z4, z5, z6 = pomoc_pomiary4['text'], pomoc_pomiary5['text'], pomoc_pomiary6['text']
    z7, z8, z9 = pomoc_pomiary7['text'], pomoc_pomiary8['text'], pomoc_pomiary9['text']
    z10, z11, z12 = pomoc_pomiary10['text'], pomoc_pomiary11['text'], pomoc_pomiary12['text']
    z13, z14, z15 = pomoc_pomiary13['text'], pomoc_pomiary14['text'], pomoc_pomiary15['text']
    z16, z17, z18 = pomoc_pomiary16['text'], pomoc_pomiary17['text'], pomoc_pomiary18['text']
    z19, z20, z21 = pomoc_pomiary19['text'], pomoc_pomiary20['text'], pomoc_pomiary21['text']
    z22, z23, z24 = pomoc_pomiary22['text'], pomoc_pomiary23['text'], pomoc_pomiary24['text']
    z25, z26, z27 = pomoc_pomiary25['text'], pomoc_pomiary26['text'], pomoc_pomiary27['text']
    z28 = pomoc_pomiary28['text']
    status = 'nierozliczone'
    stawka_vat = stawka_vat_value.get()

    nowy = DaneFormularza(data, miasto, ulica, numer, start_godzina, start_minuta, koniec_godzina, koniec_minuta,
                          stawka, p1, p2, p3, p4, p5, p6, p7, p8, p9, p10, p11, p12, p13, p14, p15, p16, p17, p18,
                          p19, p20, p21, p22, p23, p24, p25, p26, p27, p28, z1, z2, z3, z4, z5, z6, z7, z8, z9, z10,
                          z11, z12, z13, z14, z15, z16, z17, z18, z19, z20, z21, z22, z23, z24, z25, z26, z27, z28,
                          status, stawka_vat)

    dopisz_do_bazy('baza.csv', nowy)
    nowa_linia('baza.csv')
    drukuj()
    koncowy_wynik.configure(text=f'Dane wczytane poprawnie', fg='black')


# DEF DO PANELU MENU_PIERWSZY - POMIARY
def zapisz_dane():
    global kolejny_wejscie, kolejny_wejscie1, kolejny_wejscie2, kolejny_wejscie3, kolejny_wejscie4
    global kolejny_wejscie5, kolejny_wejscie6, kolejny_wejscie7, kolejny_wejscie8, kolejny_wejscie9, kolejny_wejscie10
    global kolejny_wejscie11, kolejny_wejscie12, kolejny_wejscie13, kolejny_wejscie14, kolejny_wejscie15
    global kolejny_wejscie16, kolejny_wejscie17, kolejny_wejscie18, kolejny_wejscie19, kolejny_wejscie20
    global kolejny_wejscie21, kolejny_wejscie22, kolejny_wejscie23, kolejny_wejscie24, kolejny_wejscie25
    global pierwszy_pomiar_wejscie, pierwszy_pomiar_wejscie1, root_b
    global pomoc_pomiary1, pomoc_pomiary2, pomoc_pomiary3, pomoc_pomiary4, pomoc_pomiary5, pomoc_pomiary6
    global pomoc_pomiary7, pomoc_pomiary8, pomoc_pomiary9, pomoc_pomiary10, pomoc_pomiary11, pomoc_pomiary12
    global pomoc_pomiary13, pomoc_pomiary14, pomoc_pomiary15, pomoc_pomiary16, pomoc_pomiary17, pomoc_pomiary18
    global pomoc_pomiary19, pomoc_pomiary20, pomoc_pomiary21, pomoc_pomiary22, pomoc_pomiary23, pomoc_pomiary24
    global pomoc_pomiary25, pomoc_pomiary26, pomoc_pomiary27, pomoc_pomiary28

    z1, z2, z3, z4, z5, z6, z7, z8, z9, z10, z11 = None, None, None, None, None, None, None, None, None, None, None
    z12, z13, z14, z15, z16, z17, z18, z19, z20 = None, None, None, None, None, None, None, None, None
    z21, z22, z23, z24, z25, z26, z27, z28 = None, None, None, None, None, None, None, None

    lista_z_pomiarami = [(z1, pierwszy_pomiar_wejscie, pomoc_pomiary1), (z2, pierwszy_pomiar_wejscie1, pomoc_pomiary2),
                   (z3, kolejny_wejscie, pomoc_pomiary3), (z4, kolejny_wejscie1, pomoc_pomiary4),
                   (z5, kolejny_wejscie2, pomoc_pomiary5), (z6, kolejny_wejscie3, pomoc_pomiary6),
                   (z7, kolejny_wejscie4, pomoc_pomiary7), (z8, kolejny_wejscie5, pomoc_pomiary8),
                   (z9, kolejny_wejscie6, pomoc_pomiary9), (z10, kolejny_wejscie7, pomoc_pomiary10),
                   (z11, kolejny_wejscie8, pomoc_pomiary11), (z12, kolejny_wejscie9, pomoc_pomiary12),
                   (z13, kolejny_wejscie10, pomoc_pomiary13), (z14, kolejny_wejscie11, pomoc_pomiary14),
                   (z15, kolejny_wejscie12, pomoc_pomiary15), (z16, kolejny_wejscie13, pomoc_pomiary16),
                   (z17, kolejny_wejscie14, pomoc_pomiary17), (z18, kolejny_wejscie15, pomoc_pomiary18),
                   (z19, kolejny_wejscie16, pomoc_pomiary19), (z20, kolejny_wejscie17, pomoc_pomiary20),
                   (z21, kolejny_wejscie18, pomoc_pomiary21), (z22, kolejny_wejscie19, pomoc_pomiary22),
                   (z23, kolejny_wejscie20, pomoc_pomiary23), (z24, kolejny_wejscie21, pomoc_pomiary24),
                   (z25, kolejny_wejscie22, pomoc_pomiary25), (z26, kolejny_wejscie23, pomoc_pomiary26),
                   (z27, kolejny_wejscie24, pomoc_pomiary27), (z28, kolejny_wejscie25, pomoc_pomiary28)
                   ]

    for i, element, wyjscie in lista_z_pomiarami:
        i = element.get()
        if i == '':
            wyjscie.configure(text='0')
        else:
            wyjscie.configure(text=i)

    koncowy_wynik2.configure(text='Dane wczytane poprawnie')
    root_b.destroy()


def dodaj_kolejny_pomiar():
    global liczydlo, pomiar_kolejny, pomiar_kolejny1, pomiar_kolejny2, pomiar_kolejny3, pomiar_kolejny4
    global pomiar_kolejny5, pomiar_kolejny6, pomiar_kolejny7, pomiar_kolejny8, pomiar_kolejny9, pomiar_kolejny10
    global pomiar_kolejny11, pomiar_kolejny12, kolejny_wejscie, kolejny_wejscie1, kolejny_wejscie2, kolejny_wejscie3
    global kolejny_wejscie4, kolejny_wejscie5, kolejny_wejscie6, kolejny_wejscie7, kolejny_wejscie8, kolejny_wejscie9
    global kolejny_wejscie10, kolejny_wejscie11, kolejny_wejscie12, kolejny_wejscie13, kolejny_wejscie14
    global kolejny_wejscie15, kolejny_wejscie16, kolejny_wejscie17, kolejny_wejscie18, kolejny_wejscie19
    global kolejny_wejscie20, kolejny_wejscie21, kolejny_wejscie22, kolejny_wejscie23, kolejny_wejscie24
    global kolejny_wejscie25

    licznik = int(liczydlo.cget('text'))

    pomiar = 'ZK - '
    pomiar = pomiar + str(licznik)

    if licznik == 2:
        pomiar_kolejny.configure(text=pomiar, fg=kol_czcionki, bg=kol_tla2, relief=wyglad, anchor='center',
                                 font=(czcionka, wiel_czcionki2, efekt_czcionki))
        kolejny_wejscie.place(relx=0.25, rely=0.3, relwidth=0.2, relheight=0.05)
        kolejny_wejscie1.place(relx=0.45, rely=0.3, relwidth=0.2, relheight=0.05)
    elif licznik == 3:
        pomiar_kolejny1.configure(text=pomiar, fg=kol_czcionki, bg=kol_tla2, relief=wyglad, anchor='center',
                                  font=(czcionka, wiel_czcionki2, efekt_czcionki))
        kolejny_wejscie2.place(relx=0.25, rely=0.35, relwidth=0.2, relheight=0.05)
        kolejny_wejscie3.place(relx=0.45, rely=0.35, relwidth=0.2, relheight=0.05)
    elif licznik == 4:
        pomiar_kolejny2.configure(text=pomiar, fg=kol_czcionki, bg=kol_tla2, relief=wyglad, anchor='center',
                                  font=(czcionka, wiel_czcionki2, efekt_czcionki))
        kolejny_wejscie4.place(relx=0.25, rely=0.4, relwidth=0.2, relheight=0.05)
        kolejny_wejscie5.place(relx=0.45, rely=0.4, relwidth=0.2, relheight=0.05)
    elif licznik == 5:
        pomiar_kolejny3.configure(text=pomiar, fg=kol_czcionki, bg=kol_tla2, relief=wyglad, anchor='center',
                                  font=(czcionka, wiel_czcionki2, efekt_czcionki))
        kolejny_wejscie6.place(relx=0.25, rely=0.45, relwidth=0.2, relheight=0.05)
        kolejny_wejscie7.place(relx=0.45, rely=0.45, relwidth=0.2, relheight=0.05)
    elif licznik == 6:
        pomiar_kolejny4.configure(text=pomiar, fg=kol_czcionki, bg=kol_tla2, relief=wyglad, anchor='center',
                                  font=(czcionka, wiel_czcionki2, efekt_czcionki))
        kolejny_wejscie8.place(relx=0.25, rely=0.5, relwidth=0.2, relheight=0.05)
        kolejny_wejscie9.place(relx=0.45, rely=0.5, relwidth=0.2, relheight=0.05)
    elif licznik == 7:
        pomiar_kolejny5.configure(text=pomiar, fg=kol_czcionki, bg=kol_tla2, relief=wyglad, anchor='center',
                                  font=(czcionka, wiel_czcionki2, efekt_czcionki))
        kolejny_wejscie10.place(relx=0.25, rely=0.55, relwidth=0.2, relheight=0.05)
        kolejny_wejscie11.place(relx=0.45, rely=0.55, relwidth=0.2, relheight=0.05)
    elif licznik == 8:
        pomiar_kolejny6.configure(text=pomiar, fg=kol_czcionki, bg=kol_tla2, relief=wyglad, anchor='center',
                                  font=(czcionka, wiel_czcionki2, efekt_czcionki))
        kolejny_wejscie12.place(relx=0.25, rely=0.6, relwidth=0.2, relheight=0.05)
        kolejny_wejscie13.place(relx=0.45, rely=0.6, relwidth=0.2, relheight=0.05)
    elif licznik == 9:
        pomiar_kolejny7.configure(text=pomiar, fg=kol_czcionki, bg=kol_tla2, relief=wyglad, anchor='center',
                                  font=(czcionka, wiel_czcionki2, efekt_czcionki))
        kolejny_wejscie14.place(relx=0.25, rely=0.65, relwidth=0.2, relheight=0.05)
        kolejny_wejscie15.place(relx=0.45, rely=0.65, relwidth=0.2, relheight=0.05)
    elif licznik == 10:
        pomiar_kolejny8.configure(text=pomiar, fg=kol_czcionki, bg=kol_tla2, relief=wyglad, anchor='center',
                                  font=(czcionka, wiel_czcionki2, efekt_czcionki))
        kolejny_wejscie16.place(relx=0.25, rely=0.7, relwidth=0.2, relheight=0.05)
        kolejny_wejscie17.place(relx=0.45, rely=0.7, relwidth=0.2, relheight=0.05)
    elif licznik == 11:
        pomiar_kolejny9.configure(text=pomiar, fg=kol_czcionki, bg=kol_tla2, relief=wyglad, anchor='center',
                                  font=(czcionka, wiel_czcionki2, efekt_czcionki))
        kolejny_wejscie18.place(relx=0.25, rely=0.75, relwidth=0.2, relheight=0.05)
        kolejny_wejscie19.place(relx=0.45, rely=0.75, relwidth=0.2, relheight=0.05)
    elif licznik == 12:
        pomiar_kolejny10.configure(text=pomiar, fg=kol_czcionki, bg=kol_tla2, relief=wyglad, anchor='center',
                                   font=(czcionka, wiel_czcionki2, efekt_czcionki))
        kolejny_wejscie20.place(relx=0.25, rely=0.8, relwidth=0.2, relheight=0.05)
        kolejny_wejscie21.place(relx=0.45, rely=0.8, relwidth=0.2, relheight=0.05)
    elif licznik == 13:
        pomiar_kolejny11.configure(text=pomiar, fg=kol_czcionki, bg=kol_tla2, relief=wyglad, anchor='center',
                                   font=(czcionka, wiel_czcionki2, efekt_czcionki))
        kolejny_wejscie22.place(relx=0.25, rely=0.85, relwidth=0.2, relheight=0.05)
        kolejny_wejscie23.place(relx=0.45, rely=0.85, relwidth=0.2, relheight=0.05)
    elif licznik == 14:
        pomiar_kolejny12.configure(text=pomiar, fg=kol_czcionki, bg=kol_tla2, relief=wyglad, anchor='center',
                                   font=(czcionka, wiel_czcionki2, efekt_czcionki))
        kolejny_wejscie24.place(relx=0.25, rely=0.9, relwidth=0.2, relheight=0.05)
        kolejny_wejscie25.place(relx=0.45, rely=0.9, relwidth=0.2, relheight=0.05)
    else:
        print('Wystąpił błąd')

    dodane = licznik + 1
    liczydlo.configure(text=dodane)


def pobierz_pomiary():
    global liczydlo, pierwszy_pomiar_wejscie, pierwszy_pomiar_wejscie1, wszystkie_pomiary
    global pomiar_kolejny, pomiar_kolejny1, pomiar_kolejny2, pomiar_kolejny3, pomiar_kolejny4
    global pomiar_kolejny5, pomiar_kolejny6, pomiar_kolejny7, pomiar_kolejny8, pomiar_kolejny9, pomiar_kolejny10
    global pomiar_kolejny11, pomiar_kolejny12, kolejny_wejscie, kolejny_wejscie1, kolejny_wejscie2, kolejny_wejscie3
    global kolejny_wejscie4, kolejny_wejscie5, kolejny_wejscie6, kolejny_wejscie7, kolejny_wejscie8, kolejny_wejscie9
    global kolejny_wejscie10, kolejny_wejscie11, kolejny_wejscie12, kolejny_wejscie13, kolejny_wejscie14
    global kolejny_wejscie15, kolejny_wejscie16, kolejny_wejscie17, kolejny_wejscie18, kolejny_wejscie19
    global kolejny_wejscie20, kolejny_wejscie21, kolejny_wejscie22, kolejny_wejscie23, kolejny_wejscie24
    global kolejny_wejscie25, tabela_pomiary, kolejny_wejcie, kolejny_wejscie1, kolejny_wejscie2, kolejny_wejscie3
    global kolejny_wejscie4, kolejny_wejscie5, kolejny_wejscie6, kolejny_wejscie7, kolejny_wejscie8, kolejny_wejscie9
    global kolejny_wejscie10, kolejny_wejscie11, kolejny_wejscie12, kolejny_wejscie13, kolejny_wejscie14
    global kolejny_wejscie15, kolejny_wejscie16, kolejny_wejscie17, kolejny_wejscie18, kolejny_wejscie19
    global kolejny_wejscie20, kolejny_wejscie21, kolejny_wejscie22, kolejny_wejscie23, kolejny_wejscie24
    global kolejny_wejscie25, tabela_pomocnicza_pomiary, koncowy_wynik2, pierwszy_pomiar_wejscie
    global pierwszy_pomiar_wejscie1, root_b

    root_b = tkinter.Toplevel()
    root_b.title('DragApp')
    root_b.geometry('1200x900')
    app_b = tkinter.Frame(root_b, bg='white')
    app_b.place(relx=0.025, rely=0.025, relwidth=0.95, relheight=0.95)

    logo = tkinter.PhotoImage(file='grafika\\logo.png')
    foto = tkinter.Label(master=app_b, image=logo, bg='white')
    foto.place(relx=0.7, rely=0.1, relwidth=0.3, relheight=0.2)

    label = tkinter.Label(master=app_b, text='TABELA POMIARÓW', fg=kol_czcionki, bg='red', relief=wyglad,
                          anchor='center', font=(czcionka, wiel_czcionki, efekt_czcionki))
    label.place(relx=0.05, rely=0.1, relwidth=0.6, relheight=0.05)
    label1 = tkinter.Label(master=app_b, text='PUNKT KONTROLNY', fg=kol_czcionki, bg='#37474F', relief=wyglad,
                           anchor='center',
                           font=(czcionka, wiel_czcionki, efekt_czcionki))
    label1.place(relx=0.05, rely=0.2, relwidth=0.2, relheight=0.05)
    label12 = tkinter.Label(master=app_b, text='WYNIK POMIARU', fg=kol_czcionki, bg='#37474F', relief=wyglad,
                            anchor='center',
                            font=(czcionka, wiel_czcionki, efekt_czcionki))
    label12.place(relx=0.25, rely=0.2, relwidth=0.2, relheight=0.05)
    label13 = tkinter.Label(master=app_b, text='ILOŚĆ UZIOMÓW', fg=kol_czcionki, bg='#37474F', relief=wyglad,
                            anchor='center',
                            font=(czcionka, wiel_czcionki, efekt_czcionki))
    label13.place(relx=0.45, rely=0.2, relwidth=0.2, relheight=0.05)
    pierwszy_pomiar = tkinter.Label(master=app_b, text='ZK - 1', fg=kol_czcionki, bg=kol_tla2, relief=wyglad,
                                    anchor='center', font=(czcionka, wiel_czcionki2, efekt_czcionki))
    pierwszy_pomiar.place(relx=0.05, rely=0.25, relwidth=0.2, relheight=0.05)

    pierwszy_pomiar_wejscie = tkinter.Entry(app_b, bg='white', width=30, font=('Gill sans MT', 14), justify='center')
    pierwszy_pomiar_wejscie.place(relx=0.25, rely=0.25, relwidth=0.2, relheight=0.05)
    pierwszy_pomiar_wejscie1 = tkinter.Entry(app_b, bg='white', width=30, font=('Gill sans MT', 14), justify='center')
    pierwszy_pomiar_wejscie1.place(relx=0.45, rely=0.25, relwidth=0.2, relheight=0.05)

    # DODAWANE KOLEJNYCH POZYCJI DO TABELI
    pomiar_kolejny = tkinter.Label(master=app_b, text='', bg='white')
    pomiar_kolejny.place(relx=0.05, rely=0.3, relwidth=0.2, relheight=0.05)
    kolejny_wejscie = tkinter.Entry(app_b, bg='white', width=30, font=('Gill sans MT', 14), justify='center')
    kolejny_wejscie.place_forget()
    kolejny_wejscie1 = tkinter.Entry(app_b, bg='white', width=30, font=('Gill sans MT', 14), justify='center')
    kolejny_wejscie1.place_forget()

    pomiar_kolejny1 = tkinter.Label(master=app_b, text='', bg='white')
    pomiar_kolejny1.place(relx=0.05, rely=0.35, relwidth=0.2, relheight=0.05)
    kolejny_wejscie2 = tkinter.Entry(app_b, bg='white', width=30, font=('Gill sans MT', 14), justify='center')
    kolejny_wejscie2.place_forget()
    kolejny_wejscie3 = tkinter.Entry(app_b, bg='white', width=30, font=('Gill sans MT', 14), justify='center')
    kolejny_wejscie3.place_forget()

    pomiar_kolejny2 = tkinter.Label(master=app_b, text='', bg='white')
    pomiar_kolejny2.place(relx=0.05, rely=0.4, relwidth=0.2, relheight=0.05)
    kolejny_wejscie4 = tkinter.Entry(app_b, bg='white', width=30, font=('Gill sans MT', 14), justify='center')
    kolejny_wejscie4.place_forget()
    kolejny_wejscie5 = tkinter.Entry(app_b, bg='white', width=30, font=('Gill sans MT', 14), justify='center')
    kolejny_wejscie5.place_forget()

    pomiar_kolejny3 = tkinter.Label(master=app_b, text='', bg='white')
    pomiar_kolejny3.place(relx=0.05, rely=0.45, relwidth=0.2, relheight=0.05)
    kolejny_wejscie6 = tkinter.Entry(app_b, bg='white', width=30, font=('Gill sans MT', 14), justify='center')
    kolejny_wejscie6.place_forget()
    kolejny_wejscie7 = tkinter.Entry(app_b, bg='white', width=30, font=('Gill sans MT', 14), justify='center')
    kolejny_wejscie7.place_forget()

    pomiar_kolejny4 = tkinter.Label(master=app_b, text='', bg='white')
    pomiar_kolejny4.place(relx=0.05, rely=0.5, relwidth=0.2, relheight=0.05)
    kolejny_wejscie8 = tkinter.Entry(app_b, bg='white', width=30, font=('Gill sans MT', 14), justify='center')
    kolejny_wejscie8.place_forget()
    kolejny_wejscie9 = tkinter.Entry(app_b, bg='white', width=30, font=('Gill sans MT', 14), justify='center')
    kolejny_wejscie9.place_forget()

    pomiar_kolejny5 = tkinter.Label(master=app_b, text='', bg='white')
    pomiar_kolejny5.place(relx=0.05, rely=0.55, relwidth=0.2, relheight=0.05)
    kolejny_wejscie10 = tkinter.Entry(app_b, bg='white', width=30, font=('Gill sans MT', 14), justify='center')
    kolejny_wejscie10.place_forget()
    kolejny_wejscie11 = tkinter.Entry(app_b, bg='white', width=30, font=('Gill sans MT', 14), justify='center')
    kolejny_wejscie11.place_forget()

    pomiar_kolejny6 = tkinter.Label(master=app_b, text='', bg='white')
    pomiar_kolejny6.place(relx=0.05, rely=0.6, relwidth=0.2, relheight=0.05)
    kolejny_wejscie12 = tkinter.Entry(app_b, bg='white', width=30, font=('Gill sans MT', 14), justify='center')
    kolejny_wejscie12.place_forget()
    kolejny_wejscie13 = tkinter.Entry(app_b, bg='white', width=30, font=('Gill sans MT', 14), justify='center')
    kolejny_wejscie13.place_forget()

    pomiar_kolejny7 = tkinter.Label(master=app_b, text='', bg='white')
    pomiar_kolejny7.place(relx=0.05, rely=0.65, relwidth=0.2, relheight=0.05)
    kolejny_wejscie14 = tkinter.Entry(app_b, bg='white', width=30, font=('Gill sans MT', 14), justify='center')
    kolejny_wejscie14.place_forget()
    kolejny_wejscie15 = tkinter.Entry(app_b, bg='white', width=30, font=('Gill sans MT', 14), justify='center')
    kolejny_wejscie15.place_forget()

    pomiar_kolejny8 = tkinter.Label(master=app_b, text='', bg='white')
    pomiar_kolejny8.place(relx=0.05, rely=0.7, relwidth=0.2, relheight=0.05)
    kolejny_wejscie16 = tkinter.Entry(app_b, bg='white', width=30, font=('Gill sans MT', 14), justify='center')
    kolejny_wejscie16.place_forget()
    kolejny_wejscie17 = tkinter.Entry(app_b, bg='white', width=30, font=('Gill sans MT', 14), justify='center')
    kolejny_wejscie17.place_forget()

    pomiar_kolejny9 = tkinter.Label(master=app_b, text='', bg='white')
    pomiar_kolejny9.place(relx=0.05, rely=0.75, relwidth=0.2, relheight=0.05)
    kolejny_wejscie18 = tkinter.Entry(app_b, bg='white', width=30, font=('Gill sans MT', 14), justify='center')
    kolejny_wejscie18.place_forget()
    kolejny_wejscie19 = tkinter.Entry(app_b, bg='white', width=30, font=('Gill sans MT', 14), justify='center')
    kolejny_wejscie19.place_forget()

    pomiar_kolejny10 = tkinter.Label(master=app_b, text='', bg='white')
    pomiar_kolejny10.place(relx=0.05, rely=0.8, relwidth=0.2, relheight=0.05)
    kolejny_wejscie20 = tkinter.Entry(app_b, bg='white', width=30, font=('Gill sans MT', 14), justify='center')
    kolejny_wejscie20.place_forget()
    kolejny_wejscie21 = tkinter.Entry(app_b, bg='white', width=30, font=('Gill sans MT', 14), justify='center')
    kolejny_wejscie21.place_forget()

    pomiar_kolejny11 = tkinter.Label(master=app_b, text='', bg='white')
    pomiar_kolejny11.place(relx=0.05, rely=0.85, relwidth=0.2, relheight=0.05)
    kolejny_wejscie22 = tkinter.Entry(app_b, bg='white', width=30, font=('Gill sans MT', 14), justify='center')
    kolejny_wejscie22.place_forget()
    kolejny_wejscie23 = tkinter.Entry(app_b, bg='white', width=30, font=('Gill sans MT', 14), justify='center')
    kolejny_wejscie23.place_forget()

    pomiar_kolejny12 = tkinter.Label(master=app_b, text='', bg='white')
    pomiar_kolejny12.place(relx=0.05, rely=0.9, relwidth=0.2, relheight=0.05)
    kolejny_wejscie24 = tkinter.Entry(app_b, bg='white', width=30, font=('Gill sans MT', 14), justify='center')
    kolejny_wejscie24.place_forget()
    kolejny_wejscie25 = tkinter.Entry(app_b, bg='white', width=30, font=('Gill sans MT', 14), justify='center')
    kolejny_wejscie25.place_forget()

    button2 = tkinter.Button(master=app_b, text="DODAJ", bg=kol_tla2, fg=kol_czcionki,
                             font=(czcionka, wiel_czcionki, efekt_czcionki), command=dodaj_kolejny_pomiar)
    button2.place(relx=0.75, rely=0.67, relwidth=0.2, relheight=0.05)
    button1 = tkinter.Button(master=app_b, text="ZAPISZ", bg='red', fg='white',
                             font=(czcionka, wiel_czcionki, efekt_czcionki), command=zapisz_dane)
    button1.place(relx=0.75, rely=0.725, relwidth=0.2, relheight=0.05)
    koncowy_wynik2 = tkinter.Label(master=app_b, text='', fg='black', bg='white',
                                   font=(czcionka, wiel_czcionki, efekt_czcionki), anchor='center')
    koncowy_wynik2.place(relx=0.75, rely=0.8, relwidth=0.2, relheight=0.05)

    # POMOCNICZE OBIEKTY:
    liczydlo = tkinter.Label(master=app_b, text='2', bg='white', fg='white')
    liczydlo.place(relx=0.8, rely=0.9, relwidth=0.2, relheight=0.05)
    tabela_pomiary = tkinter.Label(app_b, text='')
    tabela_pomiary.place_forget()

    root.mainloop()


# DEF DO PANELU MENU PIERWSZY - RYSOWANIE:
def pobierz_rysunek():
    global c, okno_rys, nowy_obrazek, draw, wybrany_kolor, szerokosc_pisaka, par_x1, par_y2, par_x1, par_y
    global prostokat_entry, prostokat_entry2, prostokat_entry3, prostokat_entry4
    global owal_entry, owal_entry2
    global uziom1_entry, uziom1_entry2, uziom1_entry3, uziom2_entry, uziom2_entry2, uziom2_entry3
    global uziom3_entry, uziom3_entry2, uziom3_entry3, uziom4_entry, uziom4_entry2, uziom4_entry3
    global antena_entry, antena_entry2, prosta_entry, prosta_entry2, prosta_entry3, prosta_entry4
    global tekst_entry, tekst_entry2, tekst_entry3

    okno_rys = tkinter.Toplevel()
    okno_rys.title('DragApp')
    okno_rys.geometry('1200x900')
    app_okno_rys = tkinter.Frame(okno_rys, bg='white')
    app_okno_rys.place(relx=0.025, rely=0.025, relwidth=0.95, relheight=0.95)

    c = Canvas(okno_rys, width=800, height=700, bg='white')
    c.place(relx=0.025, rely=0.025, relwidth=0.755, relheight=1)
    nowy_obrazek = Image.new("RGB", (1000, 700), 'white')
    draw = ImageDraw.Draw(nowy_obrazek)

    c.bind('<B1-Motion>', rysowanie)
    siatka()

    wybrany_kolor = 'black'
    szerokosc_pisaka = 3

    button1 = tkinter.Button(master=app_okno_rys, text="ZAPISZ", bg='red', fg='white',
                             font=(czcionka, wiel_czcionki, efekt_czcionki), command=zamknij_rysunek)
    button1.place(relx=0.8, rely=0.85, relwidth=0.2, relheight=0.05)
    button2 = tkinter.Button(master=app_okno_rys, text="WYCZYŚĆ", bg='grey', fg='white',
                             font=(czcionka, wiel_czcionki, efekt_czcionki), command=wyczysc)
    button2.place(relx=0.8, rely=0.795, relwidth=0.2, relheight=0.05)
    przycisk_siatka = PhotoImage(file='grafika\\siatka.png')
    przycisk_siatka.subsample(1, 1)
    siatka_xy = tkinter.Button(master=app_okno_rys, image=przycisk_siatka, bg='white', command=siatka)
    siatka_xy.place(relx=0.8, rely=0.9, relwidth=0.04, relheight=0.05)

    kolor_pisaka = IntVar()
    kolor_czarny = tkinter.Radiobutton(master=app_okno_rys, value='1', variable=kolor_pisaka,
                                       tristatevalue=0, activebackground='red', indicatoron=0, anchor='w', bg='black',
                                       command=wybrany_kolor_czarny)
    kolor_czarny.place(relx=0.825, rely=0.65, relwidth=0.05, relheight=0.05)
    kolor_czerwony = tkinter.Radiobutton(master=app_okno_rys, value='2', variable=kolor_pisaka,
                                         tristatevalue=1, activebackground='red', indicatoron=0, anchor='w', bg='red',
                                         command=wybrany_kolor_czerwony)
    kolor_czerwony.place(relx=0.875, rely=0.65, relwidth=0.05, relheight=0.05)
    kolor_zielony = tkinter.Radiobutton(master=app_okno_rys, value='3', variable=kolor_pisaka,
                                        tristatevalue=1, activebackground='red', indicatoron=0, anchor='w',
                                        bg='#7FFF00', command=wybrany_kolor_zielony)
    kolor_zielony.place(relx=0.925, rely=0.65, relwidth=0.05, relheight=0.05)
    kolor_niebieski = tkinter.Radiobutton(master=app_okno_rys, value='4', variable=kolor_pisaka,
                                          tristatevalue=1, activebackground='red', indicatoron=0, anchor='w', bg='blue',
                                          command=wybrany_kolor_niebieski)
    kolor_niebieski.place(relx=0.825, rely=0.7, relwidth=0.05, relheight=0.05)
    gumka = tkinter.Radiobutton(master=app_okno_rys, value='5', variable=kolor_pisaka, tristatevalue=1,
                                activebackground='red', indicatoron=0, anchor='w', fg='grey', bg='white',
                                command=wybrany_kolor_bialy)
    gumka.place(relx=0.875, rely=0.7, relwidth=0.05, relheight=0.05)
    kolor_zolty = tkinter.Radiobutton(master=app_okno_rys, value='6', variable=kolor_pisaka, tristatevalue=1,
                                      activebackground='red', indicatoron=0, anchor='w', bg='yellow',
                                      command=wybrany_kolor_zolty)
    kolor_zolty.place(relx=0.925, rely=0.7, relwidth=0.05, relheight=0.05)

    wspolrzedne = tkinter.Label(master=app_okno_rys, text='x1', fg='#DB7093', bg='white', anchor='center',
                                font=(czcionka, 10, efekt_czcionki))
    wspolrzedne.place(relx=0.841, rely=0.02, relwidth=0.035, relheight=0.03)
    wspolrzedne2 = tkinter.Label(master=app_okno_rys, text='y1', fg='#2E8B57', bg='white', anchor='center',
                                 font=(czcionka, 10, efekt_czcionki))
    wspolrzedne2.place(relx=0.875, rely=0.02, relwidth=0.035, relheight=0.03)
    wspolrzedne3 = tkinter.Label(master=app_okno_rys, text='x2', fg='#DB7093', bg='white', anchor='center',
                                 font=(czcionka, 10, efekt_czcionki))
    wspolrzedne3.place(relx=0.91, rely=0.02, relwidth=0.035, relheight=0.03)
    wspolrzedne4 = tkinter.Label(master=app_okno_rys, text='y2', fg='#2E8B57', bg='white', anchor='center',
                                 font=(czcionka, 10, efekt_czcionki))
    wspolrzedne4.place(relx=0.945, rely=0.02, relwidth=0.035, relheight=0.03)

    przycisk = PhotoImage(file='grafika\\prostokat.png')
    przycisk.subsample(1, 1)
    prostokat = tkinter.Button(master=app_okno_rys, image=przycisk, compound='left', bg='white',
                               command=rysuj_prostokat)
    prostokat.place(relx=0.8, rely=0.05, relwidth=0.041, relheight=0.05)
    prostokat_entry = tkinter.Entry(app_okno_rys, bg='white', font=('Gill sans MT', 10), justify='center')
    prostokat_entry.place(relx=0.841, rely=0.05, relwidth=0.035, relheight=0.05)
    prostokat_entry2 = tkinter.Entry(app_okno_rys, bg='white', font=('Gill sans MT', 10), justify='center')
    prostokat_entry2.place(relx=0.875, rely=0.05, relwidth=0.035, relheight=0.05)
    prostokat_entry3 = tkinter.Entry(app_okno_rys, bg='white', font=('Gill sans MT', 10), justify='center')
    prostokat_entry3.place(relx=0.91, rely=0.05, relwidth=0.035, relheight=0.05)
    prostokat_entry4 = tkinter.Entry(app_okno_rys, bg='white', font=('Gill sans MT', 10), justify='center')
    prostokat_entry4.place(relx=0.945, rely=0.05, relwidth=0.035, relheight=0.05)

    przycisk_owal = PhotoImage(file='grafika\\owal.png')
    przycisk_owal.subsample(1, 1)
    owal = tkinter.Button(master=app_okno_rys, image=przycisk_owal, compound='left', bg='white', command=rysuj_owal)
    owal.place(relx=0.8, rely=0.11, relwidth=0.041, relheight=0.05)
    owal_entry = tkinter.Entry(app_okno_rys, bg='white', font=('Gill sans MT', 10), justify='center')
    owal_entry.place(relx=0.841, rely=0.11, relwidth=0.035, relheight=0.05)
    owal_entry2 = tkinter.Entry(app_okno_rys, bg='white', font=('Gill sans MT', 10), justify='center')
    owal_entry2.place(relx=0.875, rely=0.11, relwidth=0.035, relheight=0.05)

    przycisk_uziom1 = PhotoImage(file='grafika\\uziom1.png')
    przycisk_uziom1.subsample(1, 1)
    uziom1 = tkinter.Button(master=app_okno_rys, image=przycisk_uziom1, compound='left', bg='white',
                            command=rysuj_uziom1)
    uziom1.place(relx=0.8, rely=0.17, relwidth=0.041, relheight=0.05)
    uziom1_entry = tkinter.Entry(app_okno_rys, bg='white', font=('Gill sans MT', 10), justify='center')
    uziom1_entry.place(relx=0.841, rely=0.17, relwidth=0.035, relheight=0.05)
    uziom1_entry2 = tkinter.Entry(app_okno_rys, bg='white', font=('Gill sans MT', 10), justify='center')
    uziom1_entry2.place(relx=0.875, rely=0.17, relwidth=0.035, relheight=0.05)
    uziom1_label = tkinter.Label(master=app_okno_rys, text='OPIS:', fg='black', bg='white', anchor='center',
                                 font=(czcionka, 10, efekt_czcionki))
    uziom1_label.place(relx=0.91, rely=0.17, relwidth=0.035, relheight=0.05)
    uziom1_entry3 = tkinter.Entry(app_okno_rys, bg='white', font=('Gill sans MT', 10), justify='center')
    uziom1_entry3.place(relx=0.945, rely=0.17, relwidth=0.035, relheight=0.05)

    przycisk_uziom2 = PhotoImage(file='grafika\\uziom2.png')
    przycisk_uziom2.subsample(1, 1)
    uziom2 = tkinter.Button(master=app_okno_rys, image=przycisk_uziom2, compound='left', bg='white',
                            command=rysuj_uziom2)
    uziom2.place(relx=0.8, rely=0.23, relwidth=0.041, relheight=0.05)
    uziom2_entry = tkinter.Entry(app_okno_rys, bg='white', font=('Gill sans MT', 10), justify='center')
    uziom2_entry.place(relx=0.841, rely=0.23, relwidth=0.035, relheight=0.05)
    uziom2_entry2 = tkinter.Entry(app_okno_rys, bg='white', font=('Gill sans MT', 10), justify='center')
    uziom2_entry2.place(relx=0.875, rely=0.23, relwidth=0.035, relheight=0.05)
    uziom2_label = tkinter.Label(master=app_okno_rys, text='OPIS:', fg='black', bg='white', anchor='center',
                                 font=(czcionka, 10, efekt_czcionki))
    uziom2_label.place(relx=0.91, rely=0.23, relwidth=0.035, relheight=0.05)
    uziom2_entry3 = tkinter.Entry(app_okno_rys, bg='white', font=('Gill sans MT', 10), justify='center')
    uziom2_entry3.place(relx=0.945, rely=0.23, relwidth=0.035, relheight=0.05)

    przycisk_uziom3 = PhotoImage(file='grafika\\uziom3.png')
    przycisk_uziom3.subsample(1, 1)
    uziom3 = tkinter.Button(master=app_okno_rys, image=przycisk_uziom3, compound='left', bg='white',
                            command=rysuj_uziom3)
    uziom3.place(relx=0.8, rely=0.29, relwidth=0.041, relheight=0.05)
    uziom3_entry = tkinter.Entry(app_okno_rys, bg='white', font=('Gill sans MT', 10), justify='center')
    uziom3_entry.place(relx=0.841, rely=0.29, relwidth=0.035, relheight=0.05)
    uziom3_entry2 = tkinter.Entry(app_okno_rys, bg='white', font=('Gill sans MT', 10), justify='center')
    uziom3_entry2.place(relx=0.875, rely=0.29, relwidth=0.035, relheight=0.05)
    uziom3_label = tkinter.Label(master=app_okno_rys, text='OPIS:', fg='black', bg='white', anchor='center',
                                 font=(czcionka, 10, efekt_czcionki))
    uziom3_label.place(relx=0.91, rely=0.29, relwidth=0.035, relheight=0.05)
    uziom3_entry3 = tkinter.Entry(app_okno_rys, bg='white', font=('Gill sans MT', 10), justify='center')
    uziom3_entry3.place(relx=0.945, rely=0.29, relwidth=0.035, relheight=0.05)

    przycisk_uziom4 = PhotoImage(file='grafika\\uziom4.png')
    przycisk_uziom4.subsample(1, 1)
    uziom4 = tkinter.Button(master=app_okno_rys, image=przycisk_uziom4, compound='left', bg='white',
                            command=rysuj_uziom4)
    uziom4.place(relx=0.8, rely=0.35, relwidth=0.041, relheight=0.05)
    uziom4_entry = tkinter.Entry(app_okno_rys, bg='white', font=('Gill sans MT', 10), justify='center')
    uziom4_entry.place(relx=0.841, rely=0.35, relwidth=0.035, relheight=0.05)
    uziom4_entry2 = tkinter.Entry(app_okno_rys, bg='white', font=('Gill sans MT', 10), justify='center')
    uziom4_entry2.place(relx=0.875, rely=0.35, relwidth=0.035, relheight=0.05)
    uziom4_label = tkinter.Label(master=app_okno_rys, text='OPIS:', fg='black', bg='white', anchor='center',
                                 font=(czcionka, 10, efekt_czcionki))
    uziom4_label.place(relx=0.91, rely=0.35, relwidth=0.035, relheight=0.05)
    uziom4_entry3 = tkinter.Entry(app_okno_rys, bg='white', font=('Gill sans MT', 10), justify='center')
    uziom4_entry3.place(relx=0.945, rely=0.35, relwidth=0.035, relheight=0.05)

    przycisk_antena = PhotoImage(file='grafika\\antena.png')
    przycisk_antena.subsample(1, 1)
    antena = tkinter.Button(master=app_okno_rys, image=przycisk_antena, compound='left', bg='white',
                            command=rysuj_antene)
    antena.place(relx=0.8, rely=0.41, relwidth=0.041, relheight=0.05)
    antena_entry = tkinter.Entry(app_okno_rys, bg='white', font=('Gill sans MT', 10), justify='center')
    antena_entry.place(relx=0.841, rely=0.41, relwidth=0.035, relheight=0.05)
    antena_entry2 = tkinter.Entry(app_okno_rys, bg='white', font=('Gill sans MT', 10), justify='center')
    antena_entry2.place(relx=0.875, rely=0.41, relwidth=0.035, relheight=0.05)

    przycisk_prosta = PhotoImage(file='grafika\\prosta.png')
    przycisk_prosta.subsample(1, 1)
    prosta_linia = tkinter.Button(master=app_okno_rys, image=przycisk_prosta, compound='left', bg='white',
                                  command=rysuj_prosta)
    prosta_linia.place(relx=0.8, rely=0.47, relwidth=0.041, relheight=0.05)
    prosta_entry = tkinter.Entry(app_okno_rys, bg='white', font=('Gill sans MT', 10), justify='center')
    prosta_entry.place(relx=0.841, rely=0.47, relwidth=0.035, relheight=0.05)
    prosta_entry2 = tkinter.Entry(app_okno_rys, bg='white', font=('Gill sans MT', 10), justify='center')
    prosta_entry2.place(relx=0.875, rely=0.47, relwidth=0.035, relheight=0.05)
    prosta_entry3 = tkinter.Entry(app_okno_rys, bg='white', font=('Gill sans MT', 10), justify='center')
    prosta_entry3.place(relx=0.91, rely=0.47, relwidth=0.035, relheight=0.05)
    prosta_entry4 = tkinter.Entry(app_okno_rys, bg='white', font=('Gill sans MT', 10), justify='center')
    prosta_entry4.place(relx=0.945, rely=0.47, relwidth=0.035, relheight=0.05)

    przycisk_tekst = PhotoImage(file='grafika\\abc.png')
    przycisk_tekst.subsample(1, 1)
    tekst_przycisk = tkinter.Button(master=app_okno_rys, image=przycisk_tekst, compound='left', bg='white',
                                    command=pisz_tekst)
    tekst_przycisk.place(relx=0.8, rely=0.53, relwidth=0.041, relheight=0.05)
    tekst_entry = tkinter.Entry(app_okno_rys, bg='white', font=('Gill sans MT', 10), justify='center')
    tekst_entry.place(relx=0.841, rely=0.53, relwidth=0.035, relheight=0.05)
    tekst_entry2 = tkinter.Entry(app_okno_rys, bg='white', font=('Gill sans MT', 10), justify='center')
    tekst_entry2.place(relx=0.875, rely=0.53, relwidth=0.035, relheight=0.05)
    tekst_label = tkinter.Label(master=app_okno_rys, text='OPIS:', fg='black', bg='white', anchor='center',
                                font=(czcionka, 10, efekt_czcionki))
    tekst_label.place(relx=0.91, rely=0.53, relwidth=0.035, relheight=0.05)
    tekst_entry3 = tkinter.Entry(app_okno_rys, bg='white', font=('Gill sans MT', 10), justify='center')
    tekst_entry3.place(relx=0.945, rely=0.53, relwidth=0.035, relheight=0.05)

    okno_rys.mainloop()


def siatka():
    par_x1, par_y1, par_x2, par_y2, = 25, 0, 25, 700
    par_x3, par_y3, par_x4, par_y4 = 0, 25, 1000, 25

    for i in range(40):
        c.create_line(par_x1, par_y1, par_x2, par_y2, fill="#F5DEB3", width=1)
        c.create_line(par_x3, par_y3, par_x4, par_y4, fill="#F5DEB3", width=1)
        par_x1 += 25
        par_x2 += 25
        par_y3 += 25
        par_y4 += 25

    text_x = 50
    text_y = 50

    c.create_text(28, 15, text="X", fill='#DB7093', font=(czcionka, 9, efekt_czcionki))
    c.create_text(15, 28, text="Y", fill='#2E8B57', font=(czcionka, 9, efekt_czcionki))

    for i in range(20):
        c.create_text(text_x, 15, text=str(text_x), fill='#DB7093', font=(czcionka, 8, efekt_czcionki))
        text_x += 50

    for i in range(17):
        c.create_text(15, text_y, text=str(text_y), fill='#2E8B57', font=(czcionka, 8, efekt_czcionki))
        text_y += 50


def rysuj_uziom(entry1, entry2, entry3, param):
    global wybrany_kolor
    x1 = int(entry1.get()) + param[0]
    y1 = int(entry2.get()) + param[1]
    x2 = int(entry1.get()) + param[2]
    y2 = int(entry2.get()) + param[3]

    x3 = int(entry1.get()) + param[4]
    y3 = int(entry2.get()) + param[5]
    x4 = int(entry1.get()) + param[6]
    y4 = int(entry2.get()) + param[7]

    x5 = int(entry1.get()) + param[8]
    y5 = int(entry2.get()) + param[9]
    x6 = int(entry1.get()) + param[10]
    y6 = int(entry2.get()) + param[11]

    x7 = int(entry1.get()) + param[12]
    y7 = int(entry2.get()) + param[13]
    x8 = int(entry1.get()) + param[14]
    y8 = int(entry2.get()) + param[15]

    x9 = int(entry1.get()) + param[16]
    y9 = int(entry2.get()) + param[17]
    opis = entry3.get()

    c.create_line(x1, y1, x2, y2, fill=wybrany_kolor, width=3)
    c.create_line(x3, y3, x4, y4, fill=wybrany_kolor, width=3)
    c.create_line(x5, y5, x6, y6, fill=wybrany_kolor, width=3)
    c.create_line(x7, y7, x8, y8, fill=wybrany_kolor, width=3)
    c.create_text(x9, y9, text=opis, fill='black', font=(czcionka, 24, efekt_czcionki))
    draw.line([(x1, y1), (x2, y2)], fill=wybrany_kolor, width=3)
    draw.line([(x3, y3), (x4, y4)], fill=wybrany_kolor, width=3)
    draw.line([(x5, y5), (x6, y6)], fill=wybrany_kolor, width=3)
    draw.line([(x7, y7), (x8, y8)], fill=wybrany_kolor, width=3)
    font = ImageFont.truetype('Spartan.ttf', 24)
    draw.text((x9 + param[18], y9 + param[19]), text=opis, font=font, fill='black', align="center")

    siatka()


def rysuj_uziom1():
    lista_parametrow = [-20, 0, 20, 0, -10, -8, 10, -8, -5, -16, 5, -16, 0, 0, 0, 25, 0, -50, -30, 0]
    rysuj_uziom(uziom1_entry, uziom1_entry2, uziom1_entry3, lista_parametrow)
    uziom1_entry.delete(0, 'end')
    uziom1_entry2.delete(0, 'end')
    uziom1_entry3.delete(0, 'end')


def rysuj_uziom2():
    lista_parametrow = [-20, 0, 20, 0, -10, 8, 10, 8, -5, 16, 5, 16, 0, 0, 0, -25, 0, 40, -30, 0]
    rysuj_uziom(uziom2_entry, uziom2_entry2, uziom2_entry3, lista_parametrow)
    uziom2_entry.delete(0, 'end')
    uziom2_entry2.delete(0, 'end')
    uziom2_entry3.delete(0, 'end')


def rysuj_uziom3():
    lista_parametrow = [0, -20, 0, 20, 8, -10, 8, 10, 16, -5, 16, 5, -25, 0, 0, 0, 75, 0, -30, 0]
    rysuj_uziom(uziom3_entry, uziom3_entry2, uziom3_entry3, lista_parametrow)
    uziom3_entry.delete(0, 'end')
    uziom3_entry2.delete(0, 'end')
    uziom3_entry3.delete(0, 'end')


def rysuj_uziom4():
    lista_parametrow = [0, -20, 0, 20, -8, -10, -8, 10, -16, -5, -16, 5, 25, 0, 0, 0, -75, 0, -30, 0]
    rysuj_uziom(uziom4_entry, uziom4_entry2, uziom4_entry3, lista_parametrow)
    uziom4_entry.delete(0, 'end')
    uziom4_entry2.delete(0, 'end')
    uziom4_entry3.delete(0, 'end')


def rysuj_antene():
    global wybrany_kolor
    x1 = int(antena_entry.get()) + 25
    y1 = int(antena_entry2.get()) - 25
    x2 = int(antena_entry.get())
    y2 = int(antena_entry2.get())

    c.create_line(x1, y1, x2, y2, fill=wybrany_kolor, width=3)
    draw.line([(x1, y1), (x2, y2)], fill=wybrany_kolor, width=3)

    x3 = int(antena_entry.get()) - 4
    y3 = int(antena_entry2.get()) - 4
    x4 = x3 + 8
    y4 = y3 + 8

    c.create_oval(x3, y3, x4, y4, fill=wybrany_kolor, outline=wybrany_kolor, width=3)
    draw.ellipse((x3, y3, x4, y4), fill=wybrany_kolor, outline=wybrany_kolor, width=3)
    siatka()
    antena_entry.delete(0, 'end')
    antena_entry2.delete(0, 'end')


def rysuj_prosta():
    global wybrany_kolor
    x1 = int(prosta_entry.get())
    y1 = int(prosta_entry2.get())
    x2 = int(prosta_entry3.get())
    y2 = int(prosta_entry4.get())

    c.create_line(x1, y1, x2, y2, fill=wybrany_kolor, width=3)
    draw.line([(x1, y1), (x2, y2)], fill=wybrany_kolor, width=3)
    siatka()
    prosta_entry.delete(0, 'end')
    prosta_entry2.delete(0, 'end')
    prosta_entry3.delete(0, 'end')
    prosta_entry4.delete(0, 'end')


def rysuj_prostokat():
    global prostokat_entry, prostokat_entry2, prostokat_entry3, prostokat_entry4, wybrany_kolor

    x1 = int(prostokat_entry.get())
    y1 = int(prostokat_entry2.get())
    x2 = int(prostokat_entry3.get())
    y2 = int(prostokat_entry4.get())

    c.create_rectangle(x1, y1, x2, y2, outline=wybrany_kolor, fill=None, width=3)
    draw.rectangle([(x1, y1), (x2, y2)], fill=None, outline=wybrany_kolor, width=3)
    siatka()
    prostokat_entry.delete(0, 'end')
    prostokat_entry2.delete(0, 'end')
    prostokat_entry3.delete(0, 'end')
    prostokat_entry4.delete(0, 'end')


def rysuj_owal():
    global owal_entry, owal_entry2, wybrany_kolor
    x1 = int(owal_entry.get()) - 20
    y1 = int(owal_entry2.get()) - 20
    x2 = x1 + 40
    y2 = y1 + 40

    c.create_oval(x1, y1, x2, y2, fill=None, outline=wybrany_kolor, width=3)
    draw.ellipse((x1, y1, x2, y2), fill=None, outline=wybrany_kolor, width=3)
    siatka()
    owal_entry.delete(0, 'end')
    owal_entry2.delete(0, 'end')


def pisz_tekst():
    global wybrany_kolor
    x1 = int(tekst_entry.get())
    y1 = int(tekst_entry2.get())
    opis = tekst_entry3.get()
    pomocnicza_zmienna = 8 * len(opis)

    c.create_text(x1, y1, text=opis, fill=wybrany_kolor, font=(czcionka, 24, efekt_czcionki))
    font = ImageFont.truetype('Spartan.ttf', 24)
    draw.text((x1 - pomocnicza_zmienna, y1 - 10), text=opis, font=font, fill=wybrany_kolor, align="right")
    siatka()
    tekst_entry.delete(0, 'end')
    tekst_entry2.delete(0, 'end')
    tekst_entry3.delete(0, 'end')


def wybrany_kolor_czarny():
    global wybrany_kolor
    wybrany_kolor = 'black'
    return wybrany_kolor


def wybrany_kolor_czerwony():
    global wybrany_kolor
    wybrany_kolor = 'red'
    return wybrany_kolor


def wybrany_kolor_zielony():
    global wybrany_kolor
    wybrany_kolor = '#7FFF00'
    return wybrany_kolor


def wybrany_kolor_niebieski():
    global wybrany_kolor
    wybrany_kolor = 'blue'
    return wybrany_kolor


def wybrany_kolor_bialy():
    global wybrany_kolor
    wybrany_kolor = 'white'
    return wybrany_kolor


def wybrany_kolor_zolty():
    global wybrany_kolor
    wybrany_kolor = 'yellow'
    return wybrany_kolor


def wyczysc():
    global nowy_obrazek, draw
    c.delete('all')
    nowy_obrazek = Image.new("RGB", (1000, 700), 'white')
    draw = ImageDraw.Draw(nowy_obrazek)
    siatka()


def rysowanie(event):
    global c, wybrany_kolor, szerokosc_pisaka
    if wybrany_kolor == 'white':
        x1, y1, x2, y2 = (event.x - 10), (event.y), (event.x), (event.y)
        lista = [x1, y1, x2, y2]
    else:
        x1, y1, x2, y2 = (event.x - 3), (event.y), (event.x), (event.y)
        lista = [x1, y1, x2, y2]

    c.create_line(lista, fill=wybrany_kolor, width=szerokosc_pisaka)
    draw.line(lista, wybrany_kolor, width=szerokosc_pisaka)


def zamknij_rysunek():
    global c, okno_rys, x1, x2, y1, y2, nowy_obrazek, draw, filename, poleTekstowe, poleTekstowe2, poleTekstowe3, \
           poleTekstowe4
    ImageDraw.Draw(nowy_obrazek)

    data = poleTekstowe.get()
    miasto = poleTekstowe2.get()
    ulica = poleTekstowe3.get()
    numer = poleTekstowe4.get()

    filename = 'raporty_dzienne\\' + f'{data}_{miasto}_{ulica}_{numer}.jpg'
    nowy_obrazek.save(filename)
    okno_rys.destroy()


# DEF DO PANELU MENU_PIERWSZY - ZAPISUJĄCE DO PLIKU I KONWERUJĄCE PLIKI:
def dopisz_do_bazy(plik, nowy):
    with open(plik, mode='a+', encoding='CP1250') as wejscie:
        wejscie.write(f'{nowy.data};{nowy.miasto};{nowy.ulica};{nowy.numer};{nowy.start_godzina};{nowy.start_minuta};'
                      f'{nowy.koniec_godzina};{nowy.koniec_minuta};{nowy.stawka};{nowy.p1};{nowy.p2};'
                      f'{nowy.p3};{nowy.p4};{nowy.p5};{nowy.p6};{nowy.p7};{nowy.p8};{nowy.p9};{nowy.p10};{nowy.p11};'
                      f'{nowy.p12};{nowy.p13};{nowy.p14};{nowy.p15};{nowy.p16};{nowy.p17};{nowy.p18};{nowy.p19};'
                      f'{nowy.p20};{nowy.p21};{nowy.p22};{nowy.p23};{nowy.p24};{nowy.p25};{nowy.p26};{nowy.p27};'
                      f'{nowy.p28}; {nowy.z1};{nowy.z2};{nowy.z3};{nowy.z4};{nowy.z5};{nowy.z6};{nowy.z7};{nowy.z8};'
                      f'{nowy.z9};{nowy.z10};{nowy.z11};{nowy.z12};{nowy.z13};{nowy.z14};{nowy.z15};{nowy.z16};'
                      f'{nowy.z17};{nowy.z18};{nowy.z19};{nowy.z20};{nowy.z21};{nowy.z22};{nowy.z23};{nowy.z24};'
                      f'{nowy.z25};{nowy.z26};{nowy.z27};{nowy.z28}; {nowy.status}; {nowy.stawka_vat}')

    return


def nowa_linia(plik):
    with open(plik, mode='a', encoding='CP1250') as wejscie:
        wejscie.write('\n')


def convert_to_pdf(doc):
    global word
    try:
        word = client.DispatchEx("Word.Application")
        new_name = doc.replace(".docx", r".pdf")
        worddoc = word.Documents.Open(doc)
        worddoc.SaveAs(new_name, FileFormat=17)
        worddoc.Close()

    except Exception as e:
        raise e
    finally:
        word.Quit()


def drukuj():
    global poleTekstowe, poleTekstowe2, poleTekstowe3, poleTekstowe4, root, wszystkie_pomiary, pomiary
    global tabela_pomicnicza_pomiary, lista_z_pomiarami
    data = poleTekstowe.get()
    miasto = poleTekstowe2.get()
    ulica = poleTekstowe3.get()
    numer = poleTekstowe4.get()
    lista = []
    with open('baza.csv', mode='r', encoding='CP1250') as wejscie:
        for linia in wejscie:
            dane = linia.strip().split(';')
            szukany = DaneFormularza(dane[0], dane[1], dane[2], dane[3], dane[4], dane[5], dane[6], dane[7], dane[8],
                                     dane[9], dane[10], dane[11], dane[12], dane[13], dane[14], dane[15], dane[16],
                                     dane[17], dane[18], dane[19], dane[20], dane[21], dane[22], dane[23], dane[24],
                                     dane[25], dane[26], dane[27], dane[28], dane[29], dane[30], dane[31], dane[32],
                                     dane[33], dane[34], dane[35], dane[36], dane[37], dane[38], dane[39], dane[40],
                                     dane[41], dane[42], dane[43], dane[44], dane[45], dane[46], dane[47], dane[48],
                                     dane[49], dane[50], dane[51], dane[52], dane[53], dane[54], dane[55], dane[56],
                                     dane[57], dane[58], dane[59], dane[60], dane[61], dane[62], dane[63], dane[64],
                                     )

            if szukany.data == data and szukany.miasto == miasto and szukany.ulica == ulica and szukany.numer == numer:
                lista.append(szukany)
            else:
                continue

    wpisywanie_tekstu_do_worda(lista)

    plik = f'{szukany.data}_{szukany.miasto}_{szukany.ulica}_{szukany.numer}.docx'

    convert_to_pdf(r'C:\\Users\justy\\PyCharm\projekt\\DragApp\\raporty_dzienne\\' + plik)
    wb.open_new(r'C:\\Users\\justy\PyCharm\\projekt\\DragApp\\raporty_dzienne\\' +
                f'{szukany.data}_{szukany.miasto}_{szukany.ulica}_{szukany.numer}.pdf')


def wpisywanie_tekstu_do_worda(lista):
    for element in lista:
        document = docx.Document()
        document.add_picture('grafika\\logo.png', width=Inches(1.5))
        document.add_paragraph('')
        document.add_heading(F'Data wykonania zlecenia: {element.data}', 0)
        document.add_paragraph(element.miasto.upper())
        document.add_paragraph(f'UL. {element.ulica.upper()} {element.numer}')
        document.add_heading('Rozliczenie godzin:', level=0)
        document.add_paragraph(f'Godzina rozpoczęcia: {element.start_godzina}:{element.start_minuta}\n'
                               f'Godzina zakończenia: {element.koniec_godzina}:{element.koniec_minuta}')
        document.add_heading('Zużyte materiały: ', level=0)

        materialy = (('UZIOMY', element.p1), ('ZŁĄCZKI 2 ŚRUBOWE', element.p2), ('ZŁĄCZKI 4 ŚRUBOWE', element.p3),
                     ('ZŁĄCZKI 4 ŚRUBOWE - BEDNARKA', element.p4), ('ZŁĄCZKI 45° 2 ŚRUBOWE', element.p5),
                     ('ZŁĄCZE KONTROLNE', element.p6), ('PUSZKI PODTYNKOWE', element.p7),
                     ('PUSZKI GRUNTOWE', element.p8), ('T-ki ', element.p9), ('L-ki', element.p10),
                     ('GĄSIOR METALOWY', element.p11), ('GĄSIOR MALOWANY', element.p12),
                     ('NACIĄG 20 CM', element.p13), ('KOTWY 18', element.p14), ('KOTWY 20', element.p15),
                     ('USZCZELNIACZ', element.p16), ('BETONIKI', element.p17), ('KLEJ', element.p18),
                     ('DRUT ALUMINIUM', element.p19), ('DRUT STAL', element.p20), ('POLWINIT BIAŁY', element.p21),
                     ('POLWINIT CZARNY', element.p22), ('BEDNARKA 30x4 ', element.p23), ('BEDNARKA 25x4', element.p24),
                     ('SZTYCA 4m', element.p25), ('INNE', element.p26))

        lp = 1
        table = document.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'L.P.'
        hdr_cells[1].text = 'NAZWA'
        hdr_cells[2].text = 'ILOŚĆ'
        for nazwa, ilosc in materialy:
            if ilosc != '':
                row_cells = table.add_row().cells
                row_cells[0].text = str(lp)
                row_cells[1].text = nazwa
                row_cells[2].text = ilosc
                lp += 1
            else:
                continue

        document.add_heading('Pomiary: ', level=0)

        pomiary = ((element.z1, element.z2), (element.z3, element.z4), (element.z5, element.z6),
                   (element.z7, element.z8), (element.z9, element.z10), (element.z11, element.z12),
                   (element.z13, element.z14), (element.z15, element.z16), (element.z17, element.z18),
                   (element.z19, element.z20), (element.z21, element.z22), (element.z23, element.z24),
                   (element.z25, element.z26), (element.z27, element.z28))

        if pomiary[0][0] == ' ':
            document.add_paragraph('--brak--')
        else:
            lp2 = 1
            nazwa = 'ZK - '
            licznik_nazwy = 1
            table_pomiary = document.add_table(rows=1, cols=4)
            hdr_cells = table_pomiary.rows[0].cells
            hdr_cells[0].text = 'L.P.'
            hdr_cells[1].text = 'PUNKT KONTROLNY'
            hdr_cells[2].text = 'WYNIK POMIARU'
            hdr_cells[3].text = 'ILOŚĆ UZIOMÓW'
            for wynik_pom, il_uziomow in pomiary:
                if wynik_pom != ' 0' and wynik_pom != '0':
                    row_cells = table_pomiary.add_row().cells
                    row_cells[0].text = str(lp2)
                    row_cells[1].text = str(nazwa) + str(licznik_nazwy)
                    row_cells[2].text = wynik_pom
                    row_cells[3].text = il_uziomow
                    lp2 += 1
                    nazwa = 'ZK - '
                    licznik_nazwy += 1
                else:
                    continue

        document.add_heading('Rysunek obiektu: ', level=0)
        try:
            document.add_picture('raporty_dzienne\\' + f'{element.data}_{element.miasto}_{element.ulica}_'
                                                       f'{element.numer}.jpg', width=Inches(6))
        except FileNotFoundError:
            document.add_paragraph('--brak--')

        document.save('raporty_dzienne\\' + f'{element.data}_{element.miasto}_{element.ulica}_{element.numer}.docx')


# DEF POMOCNICZE DO TWORZENIA OBIEKTÓW TKINTER

def create_labels(list_of_materials, x_place_labels, y_place_labels):
    global app_a
    wiel_czcionki_min = 8
    numeration_of_labels = 0

    for element in list_of_materials:
        element = tkinter.Label(master=app_a, text=list_of_materials[numeration_of_labels], fg=kol_czcionki,
                                bg=kol_tla2, relief=wyglad, anchor=polozenie, font=(czcionka, wiel_czcionki_min,
                                                                                    efekt_czcionki))
        element.place(relx=x_place_labels, rely=y_place_labels, relwidth=0.1, relheight=0.05)
        numeration_of_labels += 1
        x_place_labels += 0.15


def tworz_checkbutton_do_rozliczen(lista_do_rozliczenia, lista_checkboxow, y_place_labels, width, tlo, kolor_czcionki):
    global app_drugi

    licznik = 0
    dlugosc_listy = len(lista_do_rozliczenia)

    if len(lista_do_rozliczenia) <= 14:
        while licznik < dlugosc_listy:
            lista_checkboxow[licznik].configure(text=lista_do_rozliczenia[licznik], anchor='w', activebackground='red',
                                                bg=tlo, indicatoron=0, fg=kolor_czcionki, font=(czcionka, wiel_czcionki,
                                                                                                efekt_czcionki))
            lista_checkboxow[licznik].place(relx=0.1, rely=y_place_labels, relwidth=width, relheight=0.05)
            licznik += 1
            y_place_labels += 0.05
    else:
        tkinter.messagebox.showinfo(title=None,
                                    message='Wygenerowano zbyt dużo rekordów (powyżej 14).\nPodziel raport.')


def wyczysc_checkboxy(lista_checkboxow):
    global app_drugi
    licznik = 0
    for element in lista_checkboxow:
        element.configure(text='')
        element.place(relx=0, rely=0, relwidth=0, relheight=0)
        licznik += 1



# DEF DO PANELU MENU_DRUGI - WYSZUKIWANIE, POBIERANIE I ZAPISYWANIE DANYCH
def zatwierdz_przedzial_dat():
    global do_rozliczenia1, do_rozliczenia2, do_rozliczenia3, do_rozliczenia4, do_rozliczenia5, do_rozliczenia6
    global do_rozliczenia7, do_rozliczenia8, do_rozliczenia9, do_rozliczenia10, do_rozliczenia11, do_rozliczenia12
    global do_rozliczenia13, do_rozliczenia14
    global lista_dat_format, lista_checkboxow, lista_rozliczonych

    lista_checkboxow = [do_rozliczenia1, do_rozliczenia2, do_rozliczenia3, do_rozliczenia4, do_rozliczenia5,
                        do_rozliczenia6, do_rozliczenia7, do_rozliczenia8, do_rozliczenia9, do_rozliczenia10,
                        do_rozliczenia11, do_rozliczenia12, do_rozliczenia13, do_rozliczenia14]
    wyczysc_checkboxy(lista_checkboxow)

    data_od = pole_data_od.get()
    data_do = pole_data_do.get()
    lista_dat = []
    if len(data_od) == 10 and len(data_do) == 10:
        data_od_dzien = int(data_od[0:2])
        data_od_miesiac = int(data_od[3:5])
        data_od_rok = int(data_od[6:10])
        data_do_dzien = int(data_do[0:2])
        data_do_miesiac = int(data_do[3:5])
        data_do_rok = int(data_do[6:10])

        if data_od_rok == data_do_rok and data_od_miesiac <= data_do_miesiac:
            if data_od_miesiac + 1 == data_do_miesiac:
                while data_od_dzien <= 31:
                    lista_dat.append(f'{data_od_dzien}.{data_od_miesiac}.{data_od_rok}')
                    data_od_dzien += 1
                data_od_dzien = 1
                while data_od_dzien <= 31 and data_od_dzien <= data_do_dzien:
                    lista_dat.append(f'{data_od_dzien}.{data_od_miesiac + 1}.{data_od_rok}')
                    data_od_dzien += 1
            elif data_od_miesiac == data_do_miesiac:
                while data_od_dzien <= 31 and data_od_dzien <= data_do_dzien:
                    lista_dat.append(f'{data_od_dzien}.{data_od_miesiac}.{data_od_rok}')
                    data_od_dzien += 1
            else:
                tkinter.messagebox.showinfo(title=None,
                                            message='Podano błędny lub zbyt duży zakres dat (max 2 miesiące).')
        else:
            tkinter.messagebox.showinfo(title=None,
                                        message='Podano błędny lub zbyt duży zakres dat (max 2 miesiące).')
    else:
        tkinter.messagebox.showinfo(title=None,
                                    message='Podano błędny format daty.\nPrawidłowy format to DD.MM.RRRR')

    lista_dat_format = []
    for element in lista_dat:
        if len(element) == 8:
            element = '0' + element[:2] + '0' + element[2:]
            lista_dat_format.append(element)
        elif len(element) == 9:
            if element[1] == '.':
                element = '0' + element
                lista_dat_format.append(element)
            elif element[4] == '.':
                element = element[:3] + '0' + element[3:]
                lista_dat_format.append(element)
        elif len(element) == 10:
            lista_dat_format.append(element)
        else:
            tkinter.messagebox.showinfo(title=None, message='Podano błędny format daty')

    lista = []
    with open('baza.csv', mode='r', encoding='CP1250') as wejscie:
        for linia in wejscie:
            dane = linia.strip().split(';')
            pozycja_do_rozliczenia = DaneFormularza(dane[0], dane[1], dane[2], dane[3], dane[4], dane[5], dane[6],
                                                    dane[7], dane[8], dane[9], dane[10], dane[11], dane[12],
                                                    dane[13],
                                                    dane[14], dane[15], dane[16], dane[17], dane[18], dane[19],
                                                    dane[20], dane[21], dane[22], dane[23], dane[24], dane[25],
                                                    dane[26], dane[27], dane[28], dane[29], dane[30], dane[31],
                                                    dane[32], dane[33], dane[34], dane[35], dane[36], dane[37],
                                                    dane[38], dane[39], dane[40], dane[41], dane[42], dane[43],
                                                    dane[44], dane[45], dane[46], dane[47], dane[48], dane[49],
                                                    dane[50], dane[51], dane[52], dane[53], dane[54], dane[55],
                                                    dane[56], dane[57], dane[58], dane[59], dane[60], dane[61],
                                                    dane[62], dane[63], dane[64], dane[65], dane[66]
                                                    )

            for element in lista_dat_format:
                if pozycja_do_rozliczenia.data == element and pozycja_do_rozliczenia.status == ' nierozliczone':
                    lista.append(pozycja_do_rozliczenia)
                else:
                    continue

    lista_do_rozliczenia = []
    lista_rozliczonych = []

    for pozycja_do_rozliczenia in lista:
        minuty = float((int(pozycja_do_rozliczenia.start_minuta) + int(pozycja_do_rozliczenia.koniec_minuta)) / 60)
        godziny = float(int(pozycja_do_rozliczenia.koniec_godzina) - int(pozycja_do_rozliczenia.start_godzina)) + \
                  float(minuty)
        kwota_netto = float(godziny * int(pozycja_do_rozliczenia.stawka))
        kwota_do_zaplaty = float(kwota_netto + ((kwota_netto * int(pozycja_do_rozliczenia.stawka_vat)) / 100))

        godziny = str(godziny)
        godziny = godziny.replace('.', ',')
        kwota_netto = str(kwota_netto)
        kwota_netto = kwota_netto.replace('.', ',')
        kwota_do_zaplaty = str(kwota_do_zaplaty)
        kwota_do_zaplaty = kwota_do_zaplaty.replace('.', ',')

        lista_do_rozliczenia.append(f'{pozycja_do_rozliczenia.data} - {pozycja_do_rozliczenia.miasto.upper()},'
                                    f' ul. {pozycja_do_rozliczenia.ulica.upper()} {pozycja_do_rozliczenia.numer}       '
                                    f' / VAT: {pozycja_do_rozliczenia.stawka_vat} % /'
                                    )

        rozliczone_data = pozycja_do_rozliczenia.data
        rozliczone_miasto = pozycja_do_rozliczenia.miasto
        rozliczone_ulica = pozycja_do_rozliczenia.ulica
        rozliczone_numer = pozycja_do_rozliczenia.numer
        rozliczone_godziny = godziny
        rozliczone_stawka = pozycja_do_rozliczenia.stawka
        rozliczone_status = 'rozliczone'
        rozliczone_kwota_netto = kwota_netto
        rozliczone_vat = pozycja_do_rozliczenia.stawka_vat
        rozliczone_kwota = kwota_do_zaplaty

        rozliczony = DaneFormularzaRozliczone(rozliczone_data, rozliczone_miasto, rozliczone_ulica,
                                              rozliczone_numer, rozliczone_godziny, rozliczone_stawka,
                                              rozliczone_status, rozliczone_kwota_netto, rozliczone_vat,
                                              rozliczone_kwota)

        lista_rozliczonych.append(rozliczony)

    tworz_checkbutton_do_rozliczen(lista_do_rozliczenia, lista_checkboxow, 0.25, 0.6, '#DCDCDC', 'black')


def var_states():
    global lista_checkboxow, lista_rozliczonych, lista_var, root_drugi
    a = 0
    zbior_rozliczonych_do_worda = []
    suma_vat8 = 0
    suma_vat23 = 0
    razem_do_zaplaty = 0
    razem_do_zaplaty_netto = 0

    for element in lista_checkboxow:
        pobrane = lista_var[a].get()
        if pobrane == 1:
            with open('rozliczone.csv', mode='a+', encoding='UTF-8') as wyjscie:
                wyjscie.write(f'{lista_rozliczonych[a].data};{lista_rozliczonych[a].miasto};'
                              f'{lista_rozliczonych[a].ulica};{lista_rozliczonych[a].numer};'
                              f'{lista_rozliczonych[a].il_godzin};{lista_rozliczonych[a].stawka};'
                              f'{lista_rozliczonych[a].status};{lista_rozliczonych[a].kwota_netto};'
                              f'{lista_rozliczonych[a].stawka_vat};{lista_rozliczonych[a].kwota}')

                nowa_linia('rozliczone.csv')

                zbior_pomocniczy = (lista_rozliczonych[a].data, lista_rozliczonych[a].miasto,
                                    lista_rozliczonych[a].ulica, lista_rozliczonych[a].numer,
                                    lista_rozliczonych[a].kwota_netto, lista_rozliczonych[a].stawka_vat,
                                    lista_rozliczonych[a].kwota)

                pomocnicza_kwota_brutto = lista_rozliczonych[a].kwota.replace(',', '.')
                pomocnicza_kwota_netto = lista_rozliczonych[a].kwota_netto.replace(',', '.')

                if int(lista_rozliczonych[a].stawka_vat) == 8:
                    lacznie_vat8 = float(float(pomocnicza_kwota_brutto) - float(pomocnicza_kwota_netto))
                    suma_vat8 += lacznie_vat8
                elif int(lista_rozliczonych[a].stawka_vat) == 23:
                    lacznie_vat23 = float(float(pomocnicza_kwota_brutto) - float(pomocnicza_kwota_netto))
                    suma_vat23 += lacznie_vat23
                else:
                    tkinter.messagebox.showinfo(title=None, message='Nie podano stawki VAT lub podano błędna wartość')

                razem_do_zaplaty += float(pomocnicza_kwota_brutto)
                razem_do_zaplaty_netto += float(pomocnicza_kwota_netto)

            zbior_rozliczonych_do_worda.append(zbior_pomocniczy)
        a += 1

    razem_do_zaplaty = round(razem_do_zaplaty, 2)
    razem_do_zaplaty = str(razem_do_zaplaty).replace('.', ',')
    razem_do_zaplaty_netto = round(razem_do_zaplaty_netto, 2)
    razem_do_zaplaty_netto = str(razem_do_zaplaty_netto).replace('.', ',')
    suma_vat8 = round(suma_vat8, 2)
    suma_vat8 = str(suma_vat8).replace('.', ',')
    suma_vat23 = round(suma_vat23, 2)
    suma_vat23 = str(suma_vat23).replace('.', ',')

    zapisywanie_raportu_zbiorczego_do_worda(zbior_rozliczonych_do_worda, suma_vat8, suma_vat23, razem_do_zaplaty,
                                            razem_do_zaplaty_netto)

    root_drugi.destroy()


def zapisywanie_raportu_zbiorczego_do_worda(zbior_rozliczonych_do_worda, suma_vat8, suma_vat23, razem_do_zaplaty,
                                            razem_do_zaplaty_netto):
    teraz = datetime.datetime.now()
    document = docx.Document()
    document.add_paragraph(f'{" " * 130} Katowice, {teraz.day}.{teraz.month}.{teraz.year}r.')
    document.add_picture('grafika\\logo.png', width=Inches(1.5))
    document.add_paragraph('')
    document.add_paragraph('')
    document.add_heading(f'Wykonane prace:', 0)
    document.add_paragraph('')

    table = document.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].paragraphs[0].add_run('DATA').bold = True
    hdr_cells[1].paragraphs[0].add_run('ADRES').bold = True
    hdr_cells[2].paragraphs[0].add_run('KWOTA \nNETTO').bold = True
    hdr_cells[3].paragraphs[0].add_run('STAWKA \nVAT').bold = True
    hdr_cells[4].paragraphs[0].add_run('KWOTA \nBRUTTO').bold = True

    for data, miasto, ulica, numer, kwota_netto, stawka_vat, kwota_brutto in zbior_rozliczonych_do_worda:
        row_cells = table.add_row().cells
        row_cells[0].text = str(data)
        row_cells[1].text = f'{str(miasto)}, {str(ulica)} {str(numer)}'
        row_cells[2].text = str(kwota_netto) + ' zł'
        row_cells[3].text = str(stawka_vat) + ' %'
        row_cells[4].text = str(kwota_brutto) + ' zł'

    document.add_heading('', 0)
    document.add_paragraph('')
    p = document.add_paragraph('Stawka VAT 8 % suma:        ')
    p.add_run(suma_vat8).bold = True
    p.add_run(' zł')
    p = document.add_paragraph(f'Stawka VAT 23 % suma:      ')
    p.add_run(suma_vat23).bold = True
    p.add_run(' zł')
    document.add_paragraph('')
    p = document.add_paragraph('Razem do zapłaty z netto:  ')
    p.add_run(razem_do_zaplaty_netto).bold = True
    p.add_run(' zł')
    p = document.add_paragraph('')
    p.add_run(f'Lącznie do zapłaty z VAT:   {razem_do_zaplaty} zł').bold = True

    teraz = datetime.datetime.now()
    document.save('raporty_zbiorcze\\' + f'raport_zbiorczy_{teraz.day}.{teraz.month}.{teraz.year}_{teraz.hour}'
                                         f'{teraz.minute}{teraz.second}.docx')


if __name__ == '__main__':
    main()
